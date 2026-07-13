# app/__init__.py
from flask import Flask, send_from_directory, jsonify, request, session
from flask_sqlalchemy import SQLAlchemy
from flask_jwt_extended import JWTManager, verify_jwt_in_request, get_jwt_identity
from flask_cors import CORS
from flask_mail import Mail
from config import config
import os

db = SQLAlchemy()
jwt = JWTManager()
mail = Mail()


def create_app(config_name=None):
    if config_name is None:
        # Безопасный дефолт: если FLASK_ENV физически не задана на сервере,
        # считаем это production, а не development. Иначе все security-
        # проверки ниже (RuntimeError для SECRET_KEY/JWT_SECRET_KEY,
        # блокировка ADMIN_MODE в _admin_mode_enabled) молча не сработают,
        # потому что обе сравнивают именно с FLASK_ENV == 'production'.
        config_name = os.environ.get('FLASK_ENV', 'production')

    app = Flask(__name__,
                static_folder=os.path.join(os.path.dirname(os.path.dirname(__file__)), 'static'),
                static_url_path='/static')

    app.config.from_object(config.get(config_name, config['default']))

    # -----------------------------------------------------------------
    # Проверка безопасности: в production SECRET_KEY и JWT_SECRET_KEY
    # обязаны быть заданы через переменные окружения. Если остались
    # небезопасные дефолтные значения из config.py — падаем сразу,
    # а не продолжаем работать со скомпрометированной безопасностью
    # сессий/токенов для всех пользователей.
    # -----------------------------------------------------------------
    if app.config.get('FLASK_ENV') == 'production':
        insecure_defaults = {
            'SECRET_KEY': 'change-this-in-production',
            'JWT_SECRET_KEY': 'jwt-secret-change-this',
        }
        for key, default_value in insecure_defaults.items():
            if app.config.get(key) == default_value:
                raise RuntimeError(
                    f"{key} must be set via environment variable in production"
                )

    db.init_app(app)
    jwt.init_app(app)
    mail.init_app(app)

    CORS(app, origins=app.config.get('ALLOWED_ORIGINS', ['*']), supports_credentials=True)

    # Создать таблицы, если их ещё нет (SQLite, free tier, миграций пока нет).
    # db.create_all() создаёт только ОТСУТСТВУЮЩИЕ таблицы — существующие
    # не трогает и не изменяет их схему. Модели импортируются явно, иначе
    # SQLAlchemy может не знать о части таблиц на момент вызова.
    with app.app_context():
        from app import models  # noqa: F401 — регистрирует все модели в metadata
        db.create_all()

    _register_blueprints(app)
    _register_legacy_routes(app)

    root_dir = os.path.dirname(os.path.dirname(__file__))

    @app.route('/')
    def index():
        return send_from_directory(root_dir, 'index.html')

    @app.route('/login')
    def login_page():
        return send_from_directory(root_dir, 'login.html')

    @app.route('/dashboard')
    def dashboard_page():
        return send_from_directory(root_dir, 'dashboard.html')

    @app.route('/admin')
    def admin_page():
        return send_from_directory(root_dir, 'admin.html')

    @app.route('/admin-login.html')
    def admin_login_page():
        # Старый admin использует этот файл
        if os.path.exists(os.path.join(root_dir, 'admin-login.html')):
            return send_from_directory(root_dir, 'admin-login.html')
        return send_from_directory(root_dir, 'admin.html')

    @app.route('/api/admin/status', methods=['GET'])
    def admin_status():
        """Статус admin панели."""
        is_admin = 'admin' in session
        return jsonify({
            'logged_in': is_admin,
            'admin': is_admin,
            'mode': 'admin' if is_admin else 'user',
        })

    @app.route('/health')
    def health():
        return {'status': 'ok', 'service': 'ResumeAI'}

    

    return app


def _extract_text_from_request():
    """Извлечь текст резюме из запроса (JSON или файл)."""
    # Попробовать JSON
    if request.is_json:
        data = request.get_json() or {}
        return data.get('resume_text', '').strip()

    # Попробовать form data с текстом
    if request.form.get('resume_text'):
        return request.form.get('resume_text').strip()

    # Попробовать загруженный файл
    file = request.files.get('file') or request.files.get('resume')
    if file:
        filename = file.filename.lower()

        # DOCX файл
        if filename.endswith('.docx'):
            try:
                from docx import Document
                import io
                doc = Document(io.BytesIO(file.read()))
                parts = []
                # Читать параграфы
                for para in doc.paragraphs:
                    if para.text.strip():
                        parts.append(para.text)
                # Читать таблицы (резюме часто хранятся в таблицах!)
                for table in doc.tables:
                    for row in table.rows:
                        row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        # убрать дубликаты ячеек
                        seen = []
                        for t in row_texts:
                            if t not in seen:
                                seen.append(t)
                        if seen:
                            parts.append(' | '.join(seen))
                text = '\n'.join(parts)
                return text.strip()
            except Exception as e:
                from flask import current_app
                current_app.logger.error(f"[_extract_text_from_request] DOCX parse failed: {e}")
                raise ValueError("Cannot read DOCX file. Please check the file is not corrupted.")

        # TXT файл
        elif filename.endswith('.txt'):
            return file.read().decode('utf-8', errors='ignore').strip()

        # PDF файл (базовая поддержка)
        elif filename.endswith('.pdf'):
            try:
                import PyPDF2
                import io
                reader = PyPDF2.PdfReader(io.BytesIO(file.read()))
                text = '\n'.join([page.extract_text() for page in reader.pages])
                return text.strip()
            except Exception:
                raise ValueError("Cannot read PDF. Please paste text manually.")

        # Попробовать прочитать как текст
        else:
            try:
                return file.read().decode('utf-8', errors='ignore').strip()
            except Exception:
                raise ValueError("Unsupported file format")

    return ''


def _get_current_user():
    """Получить пользователя из сессии или JWT токена."""
    from app.models.user import User

    # Попробовать сессию (старый frontend)
    user_id = session.get('user_id')
    if user_id:
        return User.query.get(user_id)

    # Попробовать JWT (новый frontend)
    try:
        verify_jwt_in_request()
        user_id = int(get_jwt_identity())
        return User.query.get(user_id)
    except Exception:
        return None


def _admin_mode_enabled(app):
    """
    ADMIN_MODE разрешает бесплатный доступ без логина (для тестирования).
    Из соображений безопасности он ДОЛЖЕН игнорироваться в production,
    даже если переменная окружения ADMIN_MODE=true выставлена по ошибке
    или случайно осталась после тестов. Работает только при
    FLASK_ENV != 'production' (т.е. локально или в development/testing).
    """
    if app.config.get('FLASK_ENV') == 'production':
        return False
    return app.config.get('ADMIN_MODE', False)


def _register_legacy_routes(app):
    """Совместимые маршруты для старого index.html."""

    def _do_login():
        from app.services.auth_service import AuthService
        data = request.get_json() or {}
        result = AuthService.login(data.get('email', ''), data.get('password', ''))
        if not result['success']:
            return jsonify({'success': False, 'error': result['error']}), 401

        user = result['user']
        session['user_id'] = user.id
        session['user_email'] = user.email
        session['user_name'] = user.email.split('@')[0]
        session['premium'] = user.get_plan_name() != 'free'
        session.permanent = True

        return jsonify({
            'success': True,
            'access_token': result['access_token'],
            'refresh_token': result['refresh_token'],
            'user': {
                'id': user.id,
                'email': user.email,
                'name': user.email.split('@')[0],
                'premium': session['premium'],
                'plan': user.get_plan_name(),
            }
        })

    @app.route('/api/user/login', methods=['POST'])
    def legacy_user_login():
        return _do_login()

    @app.route('/api/auth/login', methods=['POST'])
    def legacy_auth_login():
        return _do_login()

    @app.route('/api/login', methods=['POST'])
    def legacy_login():
        """Поддерживает и admin login (username/password) и user login (email/password)."""
        from flask import current_app
        data = request.get_json() or {}

        username = data.get('username', '').strip()
        password = data.get('password', '').strip()

        # Admin credentials: логин — из ADMIN_EMAIL (config.py), пароль —
        # bcrypt-хеш из ADMIN_PASSWORD_HASH. Ни хеш, ни пароль нигде в коде
        # не хранятся в открытом виде.
        admin_login = current_app.config.get('ADMIN_EMAIL', '')
        admin_password_hash = current_app.config.get('ADMIN_PASSWORD_HASH')

        if username and admin_login and username == admin_login:
            if not admin_password_hash:
                # ADMIN_PASSWORD_HASH не задан в production — админ-роут
                # недоступен, а не работает с дефолтным паролем.
                return jsonify({
                    'success': False,
                    'error': 'Admin login is not configured on this server',
                }), 403

            import bcrypt
            try:
                is_valid = bcrypt.checkpw(
                    password.encode('utf-8'),
                    admin_password_hash.encode('utf-8'),
                )
            except (ValueError, TypeError):
                # Битый/некорректный хеш в env — тоже считаем "не сконфигурировано"
                is_valid = False

            if is_valid:
                session['admin'] = 'admin'
                session['user_id'] = None
                session['user_name'] = 'admin'
                session.permanent = True
                return jsonify({
                    'success': True,
                    'message': 'Logged in successfully',
                    'is_admin': True,
                })

            return jsonify({'success': False, 'error': 'Invalid admin credentials'}), 401

        # Обычный user login
        return _do_login()

    @app.route('/api/admin/analyze', methods=['POST'])
    def legacy_admin_analyze():
        """Admin анализ без ограничений квоты."""
        if 'admin' not in session:
            return jsonify({'error': 'Unauthorized'}), 401

        try:
            from app.services.openrouter_service import OpenRouterService

            resume_text = _extract_text_from_request()
            if not resume_text or len(resume_text) < 20:
                return jsonify({'success': False, 'error': 'Resume text is too short'}), 400

            # Admin анализ — без квоты, без пользователя
            class FakeUser:
                id = 0
            result = OpenRouterService.analyze_resume(FakeUser(), resume_text)

            if not result['success']:
                return jsonify({'success': False, 'error': result.get('error')}), 500

            return jsonify({
                'success': True,
                'analysis': result.get('analysis', ''),
                'overall_score': result.get('overall_score', 0),
                'ats_score': result.get('ats_score', 0),
                'formatting': result.get('formatting', 0),
                'content': result.get('content', 0),
                'summary': result.get('summary', ''),
                'strengths': result.get('strengths', []),
                'improvements': result.get('improvements', []),
                'key_skills': result.get('key_skills', []),
            })

        except Exception as e:
            app.logger.error(f"[legacy_admin_analyze] failed: {e}")
            return jsonify({'success': False, 'error': 'Internal server error. Please try again.'}), 500

    @app.route('/api/admin/stats', methods=['GET'])
    def legacy_admin_stats():
        """Статистика для admin панели."""
        if 'admin' not in session:
            return jsonify({'error': 'Unauthorized'}), 401
        from app.models.user import User
        from app.models.usage_log import UsageLog
        user_count = User.query.count()
        analysis_count = UsageLog.query.filter_by(action='analysis').count()
        return jsonify({
            'success': True,
            'users': user_count,
            'analyses': analysis_count,
            'active_sessions': 1,
        })

    @app.route('/api/admin/users', methods=['GET'])
    def legacy_admin_users():
        """Список пользователей для admin."""
        if 'admin' not in session:
            return jsonify({'error': 'Unauthorized'}), 401
        from app.models.user import User
        users = User.query.all()
        return jsonify({
            'success': True,
            'users': [u.to_dict() for u in users],
        })

    def _do_register():
        from app.services.auth_service import AuthService
        data = request.get_json() or {}
        result = AuthService.register(data.get('email', ''), data.get('password', ''))
        if not result['success']:
            return jsonify({'success': False, 'error': result['error']}), 400
        user = result['user']
        session['user_id'] = user.id
        session['user_email'] = user.email
        session['user_name'] = user.email.split('@')[0]
        session['premium'] = False
        return jsonify({'success': True, 'user': user.to_dict()}), 201

    @app.route('/api/user/register', methods=['POST'])
    def legacy_user_register():
        return _do_register()

    @app.route('/api/auth/register', methods=['POST'])
    def legacy_auth_register():
        return _do_register()

    @app.route('/api/user/logout', methods=['POST'])
    def legacy_logout():
        session.clear()
        return jsonify({'success': True})

    @app.route('/api/user/me', methods=['GET'])
    def legacy_user_me():
        user = _get_current_user()
        if not user:
            return jsonify({'success': False, 'error': 'Unauthorized'}), 401
        subscription = user.get_active_subscription()
        return jsonify({
            'success': True,
            'user': {
                'id': user.id,
                'email': user.email,
                'name': user.email.split('@')[0],
                'premium': user.get_plan_name() != 'free',
                'plan': user.get_plan_name(),
            },
            'subscription': subscription.to_dict() if subscription else None,
        })

    @app.route('/api/reviews', methods=['GET'])
    def legacy_reviews():
        return jsonify({
            'success': True,
            'reviews': [
                {'author': 'Sarah M.', 'text': 'Got my dream job thanks to ResumeAI!', 'rating': 5},
                {'author': 'John D.', 'text': 'Incredible tool, highly recommend!', 'rating': 5},
                {'author': 'Anna K.', 'text': 'My resume improved dramatically.', 'rating': 5},
            ]
        })

    @app.route('/api/analyze', methods=['POST'])
    def legacy_analyze():
        from flask import current_app
        admin_mode = _admin_mode_enabled(current_app)
        is_admin = 'admin' in session

        user = _get_current_user()

        # В admin-режиме или при активной admin-сессии — пропускаем без логина
        if not user and not (admin_mode or is_admin):
            return jsonify({'success': False, 'error': 'Please log in first'}), 401

        try:
            from app.services.openrouter_service import OpenRouterService
            from app.models.usage_log import UsageLog

            # Извлечь текст из запроса (JSON или файл)
            resume_text = _extract_text_from_request()

            if not resume_text or len(resume_text) < 20:
                return jsonify({'success': False, 'error': 'Resume text is too short or empty'}), 400

            # Проверить квоту (только для обычных пользователей, не admin)
            if user and not is_admin:
                subscription = user.get_active_subscription()
                if subscription and subscription.plan_name == 'free':
                    if subscription.analysis_used >= 2:
                        return jsonify({
                            'success': False,
                            'error': 'Monthly limit reached. Upgrade to Pro.'
                        }), 403

            # Использовать FakeUser если нет реального пользователя
            class FakeUser:
                id = 0
            active_user = user if user else FakeUser()

            result = OpenRouterService.analyze_resume(active_user, resume_text)
            if not result['success']:
                return jsonify({'success': False, 'error': result.get('error')}), 500

            if user and not is_admin:
                subscription = user.get_active_subscription()
                if subscription:
                    subscription.analysis_used += 1
                    db.session.commit()
                UsageLog.log(user_id=user.id, action='analysis',
                            tokens=result.get('tokens_used', 0), status='success')

            # Определить исходный формат файла
            file = request.files.get('file')
            original_format = 'txt'
            if file and file.filename:
                fname = file.filename.lower()
                if fname.endswith('.docx'):
                    original_format = 'docx'
                elif fname.endswith('.pdf'):
                    original_format = 'pdf'
                elif fname.endswith('.doc'):
                    original_format = 'doc'

            return jsonify({
                'success': True,
                'resume_text': resume_text,          # нужен фронтенду для improve
                'original_format': original_format,  # нужен для правильного download
                'analysis': result.get('analysis', ''),
                'overall_score': result.get('overall_score', 0),
                'ats_score': result.get('ats_score', 0),
                'formatting': result.get('formatting', 0),
                'formatting_score': result.get('formatting', 0),   # frontend ждёт это имя
                'content': result.get('content', 0),
                'content_score': result.get('content', 0),          # frontend ждёт это имя
                'summary': result.get('summary', ''),
                'strengths': result.get('strengths', []),
                'improvements': result.get('improvements', []),
                'key_skills': result.get('key_skills', []),
                'keywords': result.get('key_skills', []),           # saveToHistory ждёт keywords
                'detected_language': result.get('detected_language', 'en'),
            })

        except ValueError as e:
            return jsonify({'success': False, 'error': str(e)}), 400
        except Exception as e:
            app.logger.error(f"[legacy_analyze] failed: {e}")
            return jsonify({'success': False, 'error': 'Internal server error. Please try again.'}), 500

    @app.route('/api/improve', methods=['POST'])
    def legacy_improve():
        from flask import current_app
        admin_mode = _admin_mode_enabled(current_app)
        is_admin = 'admin' in session

        user = _get_current_user()
        if not user and not (admin_mode or is_admin):
            return jsonify({'success': False, 'error': 'Please log in first'}), 401

        try:
            from app.missing_routes4 import _run_improve_pipeline

            original_bytes = None
            filename = None
            resume_text_fallback = None
            original_format = 'txt'

            file = request.files.get('file') or request.files.get('resume')
            if file:
                filename = file.filename
                original_bytes = file.read()
                fname_lower = filename.lower()
                if fname_lower.endswith('.docx'):
                    original_format = 'docx'
                elif fname_lower.endswith('.pdf'):
                    original_format = 'pdf'
                elif fname_lower.endswith('.doc'):
                    original_format = 'doc'
            else:
                data = request.get_json() or {}
                resume_text_fallback = data.get('resume_text', '').strip()
                original_format = data.get('original_format', 'txt')

            if not original_bytes and not resume_text_fallback:
                return jsonify({'success': False, 'error': 'resume_text is required'}), 400

            api_key = current_app.config.get('GROQ_API_KEY')
            result = _run_improve_pipeline(original_bytes, filename, resume_text_fallback, api_key)

            if not result.get('success'):
                return jsonify({'success': False, 'error': result.get('error')}), result.get('status', 500)

            if original_bytes:
                import base64
                session['original_docx_b64'] = base64.b64encode(original_bytes).decode('ascii')
                session['item_ids'] = result['item_ids']

            return jsonify({
                'success': True,
                'improved_resume': result['improved_resume'],
                'display_text': result['display_text'],
                'original_format': original_format,
                'original_language': result.get('detected_language', 'en'),
                'item_ids': result['item_ids'],
                'quality_report': result.get('quality_report'),
                'has_original_docx': result.get('has_original_docx', False),
            })
        except Exception as e:
            import traceback
            app.logger.error("legacy_improve failed: %s\n%s", e, traceback.format_exc())
            return jsonify({'success': False, 'error': 'Internal server error. Please try again.'}), 500

    @app.route('/api/improve/docx', methods=['POST'])
    def legacy_improve_docx():
        from flask import current_app, send_file
        admin_mode = _admin_mode_enabled(current_app)
        is_admin = 'admin' in session

        user = _get_current_user()
        if not user and not (admin_mode or is_admin):
            return jsonify({'success': False, 'error': 'Please log in first'}), 401

        try:
            from app.missing_routes4 import _apply_improved_text_to_docx
            from docx import Document
            from docx.shared import Pt
            import io, base64, json as json_lib

            original_file = request.files.get('original_file')
            improved_text = request.form.get('improved_resume') or ''

            if not improved_text:
                data = request.get_json() or {}
                improved_text = data.get('improved_resume', '')

            if not improved_text:
                return jsonify({'success': False, 'error': 'No text provided'}), 400

            item_ids_raw = request.form.get('item_ids') or ''
            if item_ids_raw:
                try:
                    item_ids = json_lib.loads(item_ids_raw)
                except Exception:
                    item_ids = []
            else:
                item_ids = session.get('item_ids', [])

            if original_file:
                buf = _apply_improved_text_to_docx(original_file.read(), improved_text, item_ids)
                return send_file(buf, as_attachment=True, download_name='improved_resume.docx',
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

            b64 = session.get('original_docx_b64')
            if b64:
                buf = _apply_improved_text_to_docx(base64.b64decode(b64), improved_text, item_ids)
                return send_file(buf, as_attachment=True, download_name='improved_resume.docx',
                    mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

            # Fallback — простой текстовый DOCX
            import re as re_lib
            doc = Document()
            doc.styles['Normal'].font.size = Pt(11)
            clean = re_lib.sub(r"###ITEM_\d+###", "", improved_text)
            for line in clean.split(chr(10)):
                line = line.strip().lstrip('#').replace('**', '').replace('*', '').strip()
                doc.add_paragraph(line if line else '')
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return send_file(buf, as_attachment=True, download_name='improved_resume.docx',
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

        except Exception as e:
            app.logger.error(f"[legacy_improve_docx] failed: {e}")
            return jsonify({'success': False, 'error': 'Internal server error. Please try again.'}), 500

    # Регистрируем дополнительные маршруты из missing_routes4 (включая /api/admin/improve)
    from app.missing_routes4 import register_missing_routes
    register_missing_routes(app, _extract_text_from_request, _get_current_user)


def _register_blueprints(app):
    from app.routes.auth import auth_bp
    app.register_blueprint(auth_bp, url_prefix='/auth')

    from app.routes.user import user_bp
    app.register_blueprint(user_bp, url_prefix='/users')

    from app.routes.subscription import subscription_bp
    app.register_blueprint(subscription_bp, url_prefix='/subscription')

    from app.routes.api_keys import api_keys_bp
    app.register_blueprint(api_keys_bp, url_prefix='/api-keys')

    from app.routes.analysis import analysis_bp
    app.register_blueprint(analysis_bp, url_prefix='/analysis')