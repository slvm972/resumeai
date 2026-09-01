# app/missing_routes4.py
import os, re, uuid, time, tempfile
import langid

# ---------------------------------------------------------------------------
# Токен-маркеры — UUID-based, LLM не может "исправить" их
# ---------------------------------------------------------------------------

def _make_token():
    """Уникальный hex-токен который LLM не воспримет как осмысленный текст."""
    return "@@@" + uuid.uuid4().hex.upper()[:12] + "@@@"

# Маркер разделитель элементов
def _make_sep():
    return "###ITEM_{}###"


# ---------------------------------------------------------------------------
# Временное хранение загруженного .docx на диске — замена
# session['original_docx_b64'] (переполнял 4KB лимит cookie-based сессии
# уже на файлах ~3KB+, браузер тихо отбрасывал куку, следующий запрос молча
# уходил в деградированный fallback без форматирования оригинала).
# В session теперь кладётся только короткий токен (session['original_docx_token']),
# сам файл — на диске в tempfile.gettempdir()/resumeai_uploads/<token>.
# ---------------------------------------------------------------------------

_TEMP_UPLOAD_DIR = os.path.join(tempfile.gettempdir(), "resumeai_uploads")


def _save_temp_upload(file_bytes):
    """Сохранить файл на диск, вернуть короткий токен для session вместо base64."""
    os.makedirs(_TEMP_UPLOAD_DIR, exist_ok=True)
    _cleanup_old_temp_uploads()
    token = uuid.uuid4().hex
    path = os.path.join(_TEMP_UPLOAD_DIR, token)
    with open(path, "wb") as f:
        f.write(file_bytes)
    return token


def _load_temp_upload(token):
    """Прочитать файл по токену из session. None, если не найден/истёк."""
    if not token:
        return None
    path = os.path.join(_TEMP_UPLOAD_DIR, token)
    if not os.path.exists(path):
        return None
    with open(path, "rb") as f:
        return f.read()


def _cleanup_old_temp_uploads(max_age_seconds=3600):
    """Удалить файлы старше часа — opportunistic garbage collection на каждую новую запись."""
    if not os.path.isdir(_TEMP_UPLOAD_DIR):
        return
    now = time.time()
    for fname in os.listdir(_TEMP_UPLOAD_DIR):
        path = os.path.join(_TEMP_UPLOAD_DIR, fname)
        try:
            if os.path.isfile(path) and (now - os.path.getmtime(path)) > max_age_seconds:
                os.remove(path)
        except OSError:
            pass


# ---------------------------------------------------------------------------
# Утилиты документа
# ---------------------------------------------------------------------------

def _extract_full_text_from_docx(file_bytes):
    from docx import Document
    import io
    doc = Document(io.BytesIO(file_bytes))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts)


_LANGID_TO_NAME = {
    'en': 'English',
    'ru': 'Russian',
    'uk': 'Ukrainian',
    'he': 'Hebrew',
    'ar': 'Arabic',
    'zh': 'Chinese',
}


def _detect_language_simple(text):
    """
    Определяет язык резюме и возвращает ПОЛНОЕ НАЗВАНИЕ языка
    (не код 'ru', а 'Russian') — это контракт функции, от него
    зависит промт "Write ONLY in {detected_lang}" в
    _run_improve_pipeline. Не менять формат возврата.

    Порядок определения:
    1. Unicode fast-path для Hebrew/Arabic — эти алфавиты
       однозначно определяются по диапазону символов, без
       обращения к langid (быстрее и надёжнее на коротких строках).
    2. Для всех остальных языков (включая Russian/Ukrainian/English/
       Chinese) — langid.classify(), код результата переводится
       в полное название через словарь _LANGID_TO_NAME.
    3. Fallback — 'English', если текст пустой, langid упал с
       исключением, или вернул код, которого нет в словаре.
    """
    sample = text[:500]

    # Шаг 1: Unicode fast-path — Hebrew/Arabic (без изменений)
    counts = {
        "Hebrew": sum(1 for c in sample if "\u05d0" <= c <= "\u05ea"),
        "Arabic": sum(1 for c in sample if "\u0600" <= c <= "\u06ff"),
    }
    best_unicode = max(counts, key=counts.get)
    if counts[best_unicode] > 5:
        return best_unicode

    # Шаг 2: langid — для Russian/Ukrainian/English/Chinese/прочих
    if not sample.strip():
        return "English"

    try:
        code, _confidence = langid.classify(sample)
    except Exception:
        return "English"

    return _LANGID_TO_NAME.get(code, "English")


def _classify_para_type(para, in_table=False):
    """
    Классифицировать параграф для целей форматированного экспорта
    (RTF/ODT/PDF, Циклы B/C/D). Возвращает одно из: 'heading',
    'bullet', 'table', 'plain'.

    Проверено на реальном docx (resume_russian_engineer.docx):
    - Заголовки секций используют встроенный стиль Word 'Heading 1'
      (bold=True в runs, но решающий признак — именно style.name,
      не bold, т.к. bold встречается и у заголовка резюме "Имя Фамилия",
      который героем не является).
    - Буллеты используют стиль 'List Paragraph' И несут <w:numPr>
      в XML параграфа (numbering properties) — оба сигнала совпадают
      на реальных данных, проверяем оба для надёжности (разные
      генераторы docx могут проставить только один из двух).
    - Ячейки таблиц в этом файле не имеют собственного стиля —
      принадлежность к таблице определяется вызывающим кодом через
      параметр in_table (мы уже внутри цикла по table.rows/cells),
      а не через XML-инспекцию параграфа.
    """
    if in_table:
        return "table"

    style_name = para.style.name if para.style is not None else ""
    if style_name.startswith("Heading") or style_name.startswith("Заголовок"):
        return "heading"

    has_numpr = bool(para._p.findall(
        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr"
    ))
    if has_numpr or style_name in ("List Paragraph", "List Bullet", "List Number"):
        return "bullet"

    return "plain"


def _extract_structured(doc):
    """
    Вернуть список {'para': <Paragraph>, 'text': str, 'type': str} для всех
    непустых элементов документа, в РЕАЛЬНОМ документном порядке
    (Цикл B, Шаг 1).

    ДО этого исправления обход шёл в два раздельных прохода —
    сначала все doc.paragraphs, потом все doc.tables — что теряло
    порядок чередования параграфов и таблиц в реальном документе
    (заголовок -> таблица "должность+даты" -> буллеты -> следующая
    таблица превращалось в "все параграфы, потом все таблицы одним
    блоком"). Теперь используется doc.iter_inner_content() —
    метод Document (унаследован от BlockItemContainer, доступен
    в python-docx>=1.1.0, подтверждено в установленной версии 1.2.0),
    который отдаёт Paragraph и Table верхнего уровня в порядке их
    реального появления в документе. Проверено на синтетическом
    документе с чередованием параграф/таблица/параграф/таблица —
    порядок сохраняется, см. отчёт Цикла B.

    Таблицы обходятся с дедупликацией по identity XML-элемента ячейки
    (cell._tc), а не по позиции (ri, ci) — это необходимо для
    merge-ячеек: python-docx возвращает один и тот же объект _Cell
    на каждой позиции сетки, которую перекрывает merge.

    ВАЖНО: все ячейки таблицы сначала материализуются в один список
    (all_cells) ДО начала дедупликации. Если ходить по table.rows/row.cells
    построчно и сразу дедуплицировать, то _Cell-обёртки предыдущей строки
    становятся недостижимы и CPython может немедленно освободить их память —
    а следующий _Cell, созданный для СОВСЕМ ДРУГОЙ ячейки, может получить
    тот же id(), что и уже "виденный" объект. Тогда реальная, уникальная
    ячейка ошибочно считается дубликатом и её текст молча теряется (баг
    подтверждён эмпирически: на таблице без единого merge пропадали
    случайные ячейки). Материализация держит все _Cell живыми одновременно,
    поэтому id() не может быть переиспользован до конца обхода таблицы.

    Поле 'type' (Цикл A) — heading/bullet/table/plain, классифицируется
    _classify_para_type(). Это НЕ меняет ai_blocks/диалог с LLM (модель
    как получала/возвращала только ###ITEM_NNN### без суффикса, так и
    продолжает) — тип используется только на этапе финальной сборки
    improved_text_for_docx в _run_improve_pipeline (суффикс :TYPE
    добавляется там же), см. Цикл A.
    """
    from docx.text.paragraph import Paragraph
    from docx.table import Table

    items = []
    for block in doc.iter_inner_content():
        if isinstance(block, Paragraph):
            if block.text.strip():
                items.append({
                    "para": block,
                    "text": block.text.strip(),
                    "type": _classify_para_type(block, in_table=False),
                })
        elif isinstance(block, Table):
            all_cells = [cell for row in block.rows for cell in row.cells]
            seen = set()
            for cell in all_cells:
                key = id(cell._tc)
                if key in seen:
                    continue
                seen.add(key)
                for para in cell.paragraphs:
                    if para.text.strip():
                        items.append({"para": para, "text": para.text.strip(), "type": "table"})
    return items


def _para_has_complex_formatting(para):
    """
    True если параграф содержит сложное форматирование внутри runs:
    разные bold/italic/color/size/hyperlinks.
    Такие параграфы не трогаем — надёжность важнее агрессивного обновления.
    """
    runs = para.runs
    if len(runs) <= 1:
        return False
    # Проверяем есть ли гиперссылки в XML
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    if para._p.findall(".//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}inline"):
        return True
    # Проверяем разнородность форматирования между runs.
    # ВАЖНО: проверяем ВСЕ runs (включая пустые) — пустой run может нести
    # underline/bold и т.п., которое будет молча потеряно при упрощённой замене.
    bolds   = set(bool(r.bold)      for r in runs)
    italics = set(bool(r.italic)    for r in runs)
    underlines = set(bool(r.underline) for r in runs)
    colors  = set(r.font.color.rgb if r.font.color and r.font.color.type else None for r in runs)
    sizes   = set(r.font.size for r in runs)
    # Если есть смешанные значения — форматирование неоднородное
    if len(bolds) > 1 or len(italics) > 1 or len(underlines) > 1:
        return True
    if len(colors) > 1 or len(sizes) > 1:
        return True
    # Если ЛЮБОЙ run имеет underline=True — считаем форматирование значимым
    # и не трогаем параграф (underline часто означает email/ссылку)
    if any(r.underline for r in runs):
        return True
    return False


def _replace_para_text(para, new_text):
    """
    Run-safe замена текста параграфа.
    Стратегия:
    1. Нет runs → para.text (стандартный путь)
    2. Один run → заменяем его текст, форматирование сохраняется
    3. Однородное форматирование (все runs одинаковые) →
       весь текст в первый run, остальные очищаем (форматирование не теряется)
    4. Неоднородное форматирование (bold+normal, hyperlinks, разные цвета) →
       НЕ ТРОГАЕМ, оставляем оригинал. Надёжность важнее обновления.
    """
    runs = para.runs
    if not runs:
        para.text = new_text
        return
    if len(runs) == 1:
        runs[0].text = new_text
        return
    # Проверяем сложность форматирования
    if _para_has_complex_formatting(para):
        # Оставляем оригинал — не рискуем потерять форматирование
        return
    # Однородное форматирование — весь текст в первый run, остальные обнуляем
    runs[0].text = new_text
    for r in runs[1:]:
        r.text = ""


# ---------------------------------------------------------------------------
# Protected Tokens — детерминированная защита ДО AI
# ---------------------------------------------------------------------------

# Паттерны защищаемых сущностей (порядок важен — более специфичные первыми)
_PROTECT_PATTERNS = [
    # URL: http/https/www и LinkedIn/GitHub/GitLab
    (r"https?://[^\s]+", "URL"),
    (r"www\.[A-Za-z0-9\-]+\.[A-Za-z]{2,}[^\s]*", "URL"),
    (r"(?:linkedin\.com|github\.com|gitlab\.com)/[^\s]+", "URL"),
    # Email
    (r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}", "EMAIL"),
    # Дата рождения дд/мм/гггг или дд.мм.гггг
    (r"\d{2}[/.]\d{2}[/.]\d{4}", "DATE_BIRTH"),
    # Диапазон дат: 2023-2020 / 1997-1991 / - 2001
    (r"(?<!\d)(?:\d{4}\s*[-\u2013]\s*\d{4}|[-\u2013]\s*\d{4}|\d{4}\s*[-\u2013])(?!\d)", "DATE_RANGE"),
    # Числа с % и суффиксами K/M/B — защищаем до ID_NUM
    (r"(?<!\d)\d+(?:[,.]\d+)?\s*(?:%|K\+?|M\+?|B\+?|k\+?|m\+?|\+)(?!\d)", "NUM_METRIC"),
    # ID: 7+ цифр подряд
    (r"\b\d{7,}\b", "ID_NUM"),
    # Сертификаты известных форматов: AZ-104, AWS-SAA-C03, CCNA и т.д.
    (r"\b(?:AWS|AZ|MS|DP|AI|SC|PL|DA|MB|MD|CKA|CKS|CKAD|CCNA|CCNP|CCIE|LPIC|RHCSA|RHCE|GCP|GKE|PCEP|PMI|PMP|ITIL|CEH|OSCP|CompTIA|Security\+|Network\+|A\+|Linux\+)\s*[-:]?\s*[A-Z0-9]{2,}(?:[-][A-Z0-9]+)*\b", "CERT"),
    # MCP ID / SP коды
    (r"[A-Z0-9]{3,}(?:\s+(?:ID|SP|MCP|No|#)\s*[A-Z0-9]+)+", "CERT_CODE"),
    # Телефон: + или скобки или израильский 05X-XXXXXXX
    (r"(?:\+\d{1,3}[\s\-]?)?\(?\d{2,4}\)?[\s\-]\d{3}[\s\-]\d{4,}", "PHONE"),
    # Спецсимвольные технологии — отдельно т.к. \b не работает со спецсимволами
    (r"(?<![A-Za-z])(?:C\+\+|C#|\.NET|ASP\.NET)(?![A-Za-z0-9])", "TECH_SPECIAL"),
    # Конкретные технологии и продукты — точный список без жадного regex
    (
        r"\b(?:"
        r"Microsoft(?:\s+Windows)?|Windows(?:\s+Server)?|Office\s+365|SharePoint|Exchange|"
        r"Azure|Active\s+Directory|"
        r"Linux|Ubuntu|Debian|CentOS|Red\s*Hat|Fedora|"
        r"Cisco|Oracle|SAP|IBM|"
        r"SQL\s+Server|PostgreSQL|MySQL|MongoDB|Redis|Elasticsearch|"
        r"React(?:\.js)?|Angular|Vue(?:\.js)?|Node(?:\.js)?|Express(?:\.js)?|"
        r"Docker|Kubernetes|Terraform|Ansible|Jenkins|"
        r"Python|JavaScript|TypeScript|PHP|Ruby|Swift|Kotlin|Golang|Rust|"
        r"ASP\.NET|"
        r"AWS|GCP|VMware|Nginx|Apache|"
        r"GitHub|GitLab|Jira|Confluence|Slack|Figma|Photoshop|Illustrator"
        r")\b",
        "TECH"
    ),
    # Латинские слова 4+ символов — только если это техническая сущность:
    # CamelCase (TypeScript, GitHub, PostgreSQL), или содержит цифру (Python3, v2).
    # Обычные английские слова (Managed, Creative, Senior) НЕ токенизируются —
    # LLM должен видеть их чтобы иметь возможность улучшить текст.
    (r"[A-Za-z]+[0-9][A-Za-z0-9]*|[a-z]+[A-Z][A-Za-z0-9]*|[A-Z][a-z]+[A-Z][A-Za-z0-9]*|[A-Z]{2,}-\d+|\b[A-Z]{2,}(?:/[A-Z]{2,})+\b", "LATIN_WORD"),
]


def _protect_text(text, store):
    """
    Заменить все защищаемые подстроки в тексте на UUID-токены.
    Используем один проход — каждая позиция обрабатывается только один раз,
    исключая наложение паттернов друг на друга.
    """
    # Собираем все совпадения со всех паттернов
    matches = []
    for pattern, kind in _PROTECT_PATTERNS:
        for m in re.finditer(pattern, text):
            matches.append((m.start(), m.end(), m.group(0)))

    if not matches:
        return text

    # Сортируем по позиции, при конфликте берём самое длинное (более специфичное)
    matches.sort(key=lambda x: (x[0], -(x[1] - x[0])))

    # Убираем перекрывающиеся совпадения — жадный алгоритм
    non_overlapping = []
    last_end = 0
    for start, end, val in matches:
        if start >= last_end:
            non_overlapping.append((start, end, val))
            last_end = end

    # Строим результат
    result = []
    last_end = 0
    for start, end, val in non_overlapping:
        result.append(text[last_end:start])
        tok = _make_token()
        store[tok] = val
        result.append(tok)
        last_end = end
    result.append(text[last_end:])

    return "".join(result)


def _restore_text(text, store):
    """Восстановить все токены обратно на оригинальные значения."""
    for tok, original in store.items():
        text = text.replace(tok, original)
    return text


# ---------------------------------------------------------------------------
# Применение улучшенного текста к DOCX
# ---------------------------------------------------------------------------

def _apply_improved_text_to_docx(original_bytes, improved_text, item_ids):
    """
    Клонировать оригинальный DOCX и заменить тексты.
    Восстановление по идентификаторам ###ITEM_001### — не по индексам.
    """
    import io, re
    from docx import Document

    doc = Document(io.BytesIO(original_bytes))
    orig_items = _extract_structured(doc)

    # Разбираем ответ AI по именованным идентификаторам.
    # Цикл A: маркер теперь может нести необязательный суффикс :TYPE
    # (###ITEM_017:BULLET###) — этому пути (DOCX) сам тип не нужен, он и
    # так восстанавливает по реальным параграфам оригинала, поэтому
    # суффикс просто поглощается некапчурящей группой и отбрасывается.
    id_to_text = {}
    parts = re.split(r"###ITEM_(\d+)(?::\w+)?###", improved_text)
    # parts: ['', '001', 'текст1', '002', 'текст2', ...]
    i = 1
    while i + 1 < len(parts):
        item_id = parts[i].zfill(3)
        text = parts[i + 1].strip()
        _pre_collapse_repr = text if '•' in text else None
        if _pre_collapse_repr is not None:
            try:
                from flask import current_app
                current_app.logger.info(
                    "[DEBUG-LEAK-BULLETS] item=%s BEFORE_collapse_repr=%r newline_count=%s",
                    item_id, _pre_collapse_repr, _pre_collapse_repr.count(chr(10)),
                )
            except Exception:
                pass
        # Настоящая причина удвоения буллетов/переносов (подтверждено
        # диагностикой [DEBUG-LEAK-BULLETS] на реальном тесте): браузер при
        # отправке improved_resume через multipart/form-data (FormData) на
        # /docx-эндпоинт автоматически нормализует \n в \r\n — это
        # стандартное поведение HTML5 form-encoding, не баг браузера.
        # python-docx же при run.text = "...\r\n..." трактует \r и \n как
        # ДВА отдельных триггера разрыва строки, создавая два <w:br/>
        # вместо одного на каждый исходный перенос. Нормализуем обратно
        # в \n ДО схлопывания, иначе \n{2,} ниже не матчит \r\n\r\n вообще
        # (в нём нет двух подряд идущих чистых \n).
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        # Защитная нормализация: схлопнуть 2+ подряд идущих переноса строки
        # в один — страхует на случай отдельного, настоящего дублирования
        # (например, если LLM всё же вернёт пустую строку между буллетами).
        text = re.sub(r"\n{2,}", "\n", text)
        id_to_text[item_id] = text
        i += 2

    # [DEBUG-LEAK-BULLETS] Временная диагностика: показать точное содержимое
    # (включая непечатаемые символы через repr) для блоков с буллетами,
    # чтобы понять, что реально приходит на этом этапе — до и после collapse.
    try:
        from flask import current_app
        for _iid, _txt in id_to_text.items():
            if '•' in _txt:
                current_app.logger.info(
                    "[DEBUG-LEAK-BULLETS] item=%s after_collapse_repr=%r newline_count=%s",
                    _iid, _txt, _txt.count(chr(10)),
                )
    except Exception:
        pass

    # Применяем — ищем каждый элемент по его ID
    for idx, item in enumerate(orig_items):
        item_id = item_ids[idx] if idx < len(item_ids) else None
        if item_id and item_id in id_to_text:
            new_text = id_to_text[item_id]
            _replace_para_text(item["para"], new_text)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Flask маршруты
# ---------------------------------------------------------------------------


def _extract_facts(text):
    """
    Извлечь из текста факты которые нельзя выдумывать:
    числа, проценты, года, слова с заглавной буквы (имена, компании).
    Возвращает множество строк-фактов.

    АРХИТЕКТУРНОЕ ПРАВИЛО: слово в самом начале текста или сразу после
    '. '/'! '/'? ' (начало нового предложения) НЕ считается фактом, даже
    если оно написано с заглавной буквы — это грамматическое требование
    английского языка (Sentence-initial capitalization), а не индикатор
    имени собственного/компании/технологии. Слова с заглавной буквы в
    середине предложения (Google, Microsoft, React) по-прежнему считаются
    фактами и ловятся.
    """
    facts = set()
    # Числа (включая проценты, суммы)
    for m in re.finditer(r"(?<!\d)\d+(?:[.,]\d+)?\s*(?:%|тыс|млн|k|K)?(?!\d)", text):
        facts.add(m.group(0).strip())

    # Позиции начала предложений: позиция 0 и позиции сразу после ". "/"! "/"? "
    sentence_start_positions = {0}
    for m in re.finditer(r"[.!?]\s+", text):
        sentence_start_positions.add(m.end())

    for m in re.finditer(r"\b[A-ZА-ЯЁ][A-Za-zА-ЯЁа-яё]{2,}\b", text):
        if m.start() in sentence_start_positions:
            continue

        facts.add(m.group(0))
    return facts


_FABRICATED_CLAIM_RE = re.compile(
    r"\bresulting in\b|\bensuring\b|\bleveraging\b|\bfostering\b|"
    r"\bstreamlining\b|\ballowing\b|\bdriving\b|\benhancing\b|"
    r"\bmeet(?:ing)? user needs\b|\bexceed(?:ing)? expectations\b|"
    r"\bstrong network\b|\bdeep understanding\b",
    re.IGNORECASE,
)


def _validate_block(orig_text, new_text):
    """
    Проверить что новый текст не содержит фактов отсутствующих в оригинале,
    а также не содержит фабрикованных причинно-следственных утверждений
    (LLM любит дописывать "resulting in...", "ensuring..." и т.п. — это
    недоказанные claims, которых не было в оригинале).
    Возвращает (is_valid: bool, reason: str).
    """
    if not new_text or not orig_text:
        return True, ""
    orig_facts = _extract_facts(orig_text)
    new_facts  = _extract_facts(new_text)
    invented   = new_facts - orig_facts
    # Фильтруем: цифры 1-4 это вероятно пункты списка, не факты
    invented = {f for f in invented if not re.match(r"^\d$", f)}
    if invented:
        return False, f"Invented facts: {', '.join(sorted(invented)[:5])}"

    fabricated = set(_FABRICATED_CLAIM_RE.findall(new_text)) - set(_FABRICATED_CLAIM_RE.findall(orig_text))
    if fabricated:
        return False, f"Fabricated claim added: {', '.join(sorted(fabricated)[:3])}"

    return True, ""


# ---------------------------------------------------------------------------
# Классификация элементов — что замораживать, что улучшать, что защищать частично
# ---------------------------------------------------------------------------

# Заголовки секций — всегда заморозка
SECTION_HEADERS_SET = {
    "ניסיון תעסוקתי",  # ניסיון תעסוקתי
    "השכלה",         # השכלה
    "שפות",               # שפות
    "יכולת מקצועית",  # יכולת מקצועית
    "כתובת",         # כתובת
    "תאריך לידה",  # תאריך לידה
    "ת.ז",                           # ת.ז
    "עברית",         # עברית
    "רוסית -",       # רוסית -
    "אנגלית",   # אנגלית
    "רוסיט",         # русский заголовок если есть
    "Опыт работы", "Образование", "Навыки", "Языки", "Контакты",
    "Experience", "Education", "Skills", "Languages", "Contacts", "Summary",
    "Professional Summary", "Work Experience", "Career Highlights",
    "Professional Experience", "Employment History", "Contact",
    "Selected Publications", "Grants and Awards", "Academic Appointments",
    "Awards",
}

# Паттерны для классификации элементов
_RE_PHONE    = re.compile(r"(?:\+\d{1,3}[\s\-]?)?\(?\d{2,4}\)?[\s\-]\d{3}[\s\-]\d{4,}|\d{9,}")
_RE_EMAIL    = re.compile(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}")
_RE_DATE_RANGE = re.compile(r"(?<!\d)(?:\d{4}\s*[-–]\s*\d{4}|[-–]\s*\d{4}|\d{4}\s*[-–])(?!\d)")
_RE_DATE_BIRTH = re.compile(r"\d{2}[/.]\d{2}[/.]\d{4}")
_RE_ID_NUM   = re.compile(r"\d{7,}")
_RE_URL      = re.compile(r"https?://|www\.")
_RE_LANG_LINE = re.compile(  # строка описания языка: "דיבור רמה גבוהה..."
    r"דיבור|קריאה|כתיבה"
)
# Identity-строка вида "Title — Company, Location" или "Degree — University" —
# структурная строка (должность/степень/место), не должна переписываться,
# как и контактный блок. Обнаружено на реальных данных: такие строки
# систематически уходили в improve и портились при retry.
_RE_IDENTITY_DASH = re.compile(r"\s[—–-]\s")
# Короткий label-заголовок таблицы навыков вида "Marketing Tools:", "Design Tools:"
_RE_COLON_LABEL = re.compile(r":\s*$")

# Слова-маркеры которые указывают что текст — описание работы (нужно улучшать)
# Слова-маркеры описаний работы — Hebrew и English глаголы
_RE_IMPROVABLE = re.compile(
    r"אחראי|פיתוח|הטמעה|ניהל|סיפק|בנייה|עבדתי|תכנון|הקמתי|פתרתי|"
    r"למידה|התקנה|תמיכה|הכשרה|הדרכה|עבודה|ייעוץ|ביצוע|שיפור|הובלה|"
    r"managed|developed|implemented|designed|built|maintained|led|created|"
    r"improved|installed|configured|supported|trained|analyzed|optimized|"
    r"saved|achieved|increased|reduced|delivered|launched|drove|grew",
    re.IGNORECASE
)


def _classify_item(text, idx, total_items):
    """
    Простая детерминированная классификация:
    - freeze: section headers, строки только из данных (дата/ID/телефон), описания языков
    - improve: всё остальное (улучшаем с токенами на факты внутри)
    Нет protect-категории — это упрощает логику и уменьшает число замороженных блоков.
    """
    t = text.strip()

    # Section headers — заморозить
    if t in SECTION_HEADERS_SET:
        return "freeze"

    # Очень короткие строки — только данные, нечего улучшать.
    # Исключение: если строка содержит глагол-маркер достижения — это
    # короткое, но реальное достижение (например "Led 5 projects",
    # "Saved $50K") и должно идти на улучшение, а не замораживаться.
    if len(t) < 15 and not _RE_IMPROVABLE.search(t):
        return "freeze"

    # Описание языков (דיבור / קריאה / כתיבה)
    if _RE_LANG_LINE.search(t):
        return "freeze"

    # Identity-строка "Title — Company, Location" / "Degree — University" —
    # структурное поле, не должно переписываться (нет глагола-маркера)
    if _RE_IDENTITY_DASH.search(t) and not _RE_IMPROVABLE.search(t):
        return "freeze"

    # Короткий label заголовок таблицы навыков ("Marketing Tools:", "Skills:")
    if _RE_COLON_LABEL.search(t) and not _RE_IMPROVABLE.search(t):
        return "freeze"

    # Перечисление имён собственных через запятую (инструменты, языки,
    # технологии) без глагола-маркера — "HubSpot, Salesforce, Marketo",
    # "English (Native), German (Conversational)". Нечего улучшать в списке
    # названий — только риск фабрикации при попытке "переписать".
    if not _RE_IMPROVABLE.search(t):
        parts = [p.strip() for p in t.split(",") if p.strip()]
        if len(parts) >= 2:
            all_short = all(len(p.split()) <= 4 for p in parts)
            all_capitalized = all(p[0].isupper() for p in parts if p)
            if all_short and all_capitalized:
                return "freeze"

    # Контактный блок (email + телефон/URL в одной строке) — структурное поле,
    # как и раздел языков, не должно переписываться в prose.
    if _RE_EMAIL.search(t) and (_RE_PHONE.search(t) or _RE_URL.search(t)):
        return "freeze"

    # Строка состоит ТОЛЬКО из данных — нет слов для улучшения
    # Убираем все найденные токены и смотрим что осталось
    store_tmp = {}
    cleaned = _protect_text(t, store_tmp)
    # После защиты остался только текст без данных
    leftover = cleaned
    for tok in store_tmp:
        leftover = leftover.replace(tok, ' ')
    leftover_words = [w for w in leftover.split() if len(w) > 2 and w not in ('ו-', 'של', 'עם', 'את', 'על')]
    if len(leftover_words) == 0 and len(t) < 40 and not _RE_IMPROVABLE.search(t):
        return "freeze"

    # Всё остальное — улучшать
    return "improve"


# ---------------------------------------------------------------------------
# Quality Gate — проверка качества улучшения перед принятием
# ---------------------------------------------------------------------------

def _text_similarity(a, b):
    """
    Простое сходство двух строк на основе общих символьных биграмм.
    Возвращает float 0.0–1.0. Не требует внешних библиотек.
    """
    def bigrams(s):
        s = s.strip().lower()
        return [s[i:i+2] for i in range(len(s)-1)] if len(s) > 1 else [s]
    bg_a = bigrams(a)
    bg_b = set(bigrams(b))
    if not bg_a or not bg_b:
        return 1.0
    matches = sum(1 for bg in bg_a if bg in bg_b)
    return matches / max(len(bg_a), len(bg_b))


def _has_quality_improvement(orig, improved):
    """
    Проверить что улучшение реально качественное:
    - не просто добавлены/убраны пробелы
    - есть реальные изменения в словах
    Returns (is_quality: bool, reason: str)
    """
    # Нормализуем для сравнения
    orig_words = set(orig.strip().split())
    impr_words = set(improved.strip().split())
    # Новые слова которых не было
    new_words = impr_words - orig_words
    # Убранные слова
    removed_words = orig_words - impr_words
    # Хоть что-то изменилось на уровне слов
    has_word_change = bool(new_words or removed_words)
    # Длина изменилась более чем на 10%
    len_change = abs(len(improved) - len(orig)) / max(len(orig), 1)
    has_length_change = len_change > 0.10
    if has_word_change or has_length_change:
        return True, f"changed_words={len(new_words)+len(removed_words)} len_delta={len_change:.1%}"
    return False, "no_word_changes"


def _has_genuine_word_change(orig, improved):
    """
    True если среди изменённых слов есть хотя бы одно, отличающееся не
    только знаками препинания (защита от ложного срабатывания когда
    'word' и 'word,' считаются как new+removed word).
    """
    strip_punct = lambda w: w.strip(".,;:!?—-")
    orig_words = {strip_punct(w) for w in orig.strip().split()}
    impr_words = {strip_punct(w) for w in improved.strip().split()}
    new_words = impr_words - orig_words
    removed_words = orig_words - impr_words
    return bool(new_words or removed_words)


def _quality_gate(orig_text, improved_text, threshold=0.95):
    """
    Quality Gate: проверить нужна ли повторная попытка.
    Returns (accepted: bool, similarity: float, reason: str)

    Порядок проверки: сначала confirmed word-level change (реальные
    изменённые/добавленные/убранные слова) — если он подтверждён, это
    ПРИНИМАЕТСЯ независимо от высокого биграммного сходства (замена
    одного сильного глагола на синоним в короткой фразе даёт sim>95%
    просто из-за общей длины окружающего текста). Только при отсутствии
    подтверждённых словесных изменений применяется порог similarity.
    """
    sim = _text_similarity(orig_text, improved_text)
    quality_ok, quality_reason = _has_quality_improvement(orig_text, improved_text)
    if quality_ok and _has_genuine_word_change(orig_text, improved_text):
        return True, sim, f"accepted similarity={sim:.3f} {quality_reason}"
    if sim > threshold:
        return False, sim, f"similarity={sim:.3f} above threshold={threshold}"
    return False, sim, f"no_quality_improvement ({quality_reason})"

def _extract_retry_after_seconds(error_message, default=2.0, cap=12.0):
    """
    Извлечь время ожидания из сообщения об ошибке Groq вида
    "Please try again in 5.67s." Возвращает секунды, ограниченные [0, cap].
    Если не удалось распарсить — возвращает default.
    """
    m = re.search(r"try again in ([\d.]+)s", error_message or "")
    if not m:
        return default
    try:
        return min(float(m.group(1)), cap)
    except ValueError:
        return default


def _run_improve_pipeline(original_bytes, filename, resume_text_fallback, api_key):
    """
    Общий защищённый pipeline улучшения резюме:
    Protected Tokens -> LLM -> Fact Validation -> Quality Gate -> Retry.

    Параметры:
      original_bytes       — байты загруженного файла (или None)
      filename              — имя файла (для определения .docx) или None
      resume_text_fallback  — текст резюме, если файла нет (JSON-путь)
      api_key               — GROQ_API_KEY

    Возвращает dict:
      {"success": True, ...}  — при успехе, те же поля что раньше отдавал /api/admin/improve
      {"success": False, "error": str, "status": int} — при ошибке
    """
    resume_text = ""
    orig_items = []
    NL = chr(10)

    is_docx = bool(filename and filename.lower().endswith(".docx"))

    if original_bytes is not None:
        if is_docx:
            resume_text = _extract_full_text_from_docx(original_bytes)
        else:
            resume_text = original_bytes.decode("utf-8", errors="ignore")
    else:
        resume_text = (resume_text_fallback or "").strip()

    if not resume_text or len(resume_text) < 20:
        return {"success": False, "error": "Resume text too short", "status": 400}

    if not api_key:
        return {"success": False, "error": "GROQ_API_KEY not configured", "status": 500}

    import requests as req_lib

    detected_lang = _detect_language_simple(resume_text)

    if original_bytes is not None and is_docx:
        from docx import Document
        import io
        doc_tmp = Document(io.BytesIO(original_bytes))
        orig_items = _extract_structured(doc_tmp)
    else:
        orig_items = [{"text": l, "type": "plain"} for l in resume_text.split(NL) if l.strip()]

    try:
        from flask import current_app
        current_app.logger.info(
            "[DEBUG-LEAK] _run_improve_pipeline: filename=%s text_preview=%r orig_items_count=%s",
            filename,
            resume_text[:80],
            len(orig_items),
        )
    except Exception:
        pass

    # -----------------------------------------------------------
    # Шаг 1: Protected Tokens — защита ДО AI
    # -----------------------------------------------------------
    store = {}       # token -> original_value
    item_ids = []    # ID каждого элемента (001, 002, ...)
    ai_blocks = []   # блоки для AI с именованными идентификаторами
    strategy_map = {}  # item_id -> реально применённая strategy (freeze/improve/protect)

    n_items = len(orig_items)
    for i, item in enumerate(orig_items):
        item_id = str(i + 1).zfill(3)
        item_ids.append(item_id)
        text = item["text"]

        # Первые 2 элемента — имя и телефон/email — всегда заморозка
        if i <= 1:
            strategy = "freeze"
        else:
            strategy = _classify_item(text, i, n_items)

        strategy_map[item_id] = strategy

        if strategy == "freeze":
            # Заморозить целиком — AI не видит содержимое
            tok = _make_token()
            store[tok] = text
            ai_blocks.append(f"###ITEM_{item_id}###\n{tok}")

        elif strategy == "improve":
            # Улучшать — токены только на факты внутри текста
            if NL in text:
                lines = text.split(NL)
                protected_lines = [_protect_text(l, store) for l in lines]
                ai_blocks.append(f"###ITEM_{item_id}###\n" + NL.join(protected_lines))
            else:
                ai_blocks.append(f"###ITEM_{item_id}###\n" + _protect_text(text, store))

        else:  # protect
            # Данные с контекстом — весь текст через токены
            tok = _make_token()
            store[tok] = text
            ai_blocks.append(f"###ITEM_{item_id}###\n{tok}")

    n = len(ai_blocks)
    ai_input = "\n\n".join(ai_blocks)

    # -----------------------------------------------------------
    # Шаг 2: Запрос к AI
    # -----------------------------------------------------------
    system_prompt = (
        f"You are a professional resume editor.\n\n"
        f"RULES:\n"
        f"1. Write ONLY in {detected_lang}\n"
        f"2. Input has {n} blocks, each starting with ###ITEM_NNN###\n"
        f"3. Return ALL {n} blocks in the SAME order with the SAME ###ITEM_NNN### identifiers\n"
        f"4. Tokens like @@@A1B2C3D4E5F6@@@ are protected values — copy them EXACTLY as-is\n"
        f"5. Improve ONLY: job descriptions and skill descriptions — use stronger, more precise action verbs for what is already described. Do not add new clauses, outcomes, or explanations.\n"
        f"6. NEVER downgrade a verb or phrase to something weaker, more generic, or less professional than the original (example of a FORBIDDEN downgrade: \"Collaborated with\" → \"Worked with\"). Only replace a word if the replacement is strictly stronger or more precise (example of a CORRECT upgrade: \"Managed\" → \"Directed\"). If you are not confident the replacement is stronger, leave the original word unchanged.\n"
        f"7. Keep unchanged: everything that is a token, section headers, dates, IDs\n"
        f"8. Multiline items: keep same number of lines, single newline between them\n"
        f"9. Do NOT merge blocks, do NOT split blocks, do NOT add extra ###ITEM### markers\n"
        f"10. NEVER invent or add anything not in the original: no new jobs, certifications, courses, achievements, responsibilities, skills, education, outcomes, results, or causal explanations (phrases like \"resulting in\", \"which improved\", \"leading to\", \"by leveraging\", \"ensuring\", \"driving\"). If a sentence has nothing to strengthen, return it unchanged rather than adding filler."
    )

    user_prompt = (
        f"Rewrite these {n} resume blocks according to the rules above. "
        f"Return with ###ITEM_NNN### identifiers.\n\n"
        f"{ai_input}\n\nOUTPUT ({n} blocks):"
    )

    payload = {
        "model": "openai/gpt-oss-120b",
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_prompt},
        ],
        "temperature": 0.15,
        "max_tokens": 4000,
    }

    resp = req_lib.post(
        "https://api.groq.com/openai/v1/chat/completions",
        headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
        json=payload, timeout=90,
    )

    if resp.status_code == 429:
        payload["model"] = "openai/gpt-oss-20b"
        resp = req_lib.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json=payload, timeout=90,
        )

    if resp.status_code == 429:
        # Оба варианта модели упёрлись в rate limit — ждём столько,
        # сколько сама Groq API просит подождать (bounded, максимум 12с),
        # и делаем один финальный повтор, прежде чем сдаться.
        err_msg = resp.json().get("error", {}).get("message", "")
        wait_s = _extract_retry_after_seconds(err_msg)
        time.sleep(wait_s)
        resp = req_lib.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json=payload, timeout=90,
        )

    if resp.status_code != 200:
        err = resp.json().get("error", {}).get("message", "Groq API error")
        return {"success": False, "error": err, "status": 500}

    raw_response = resp.json()["choices"][0]["message"]["content"]
    tokens = resp.json().get("usage", {}).get("total_tokens", 0)

    # -----------------------------------------------------------
    # Шаг 3: Восстановление + Quality Gate + Retry + Отчёт
    # -----------------------------------------------------------
    def _parse_ai_response(raw):
        """Разобрать ответ AI по идентификаторам блоков."""
        result = {}
        parts = re.split(r"###ITEM_(\d+)###", raw)
        k = 1
        while k + 1 < len(parts):
            iid = parts[k].zfill(3)
            text = parts[k + 1].strip()
            text = re.sub(NL + r"{2,}", NL, text)
            result[iid] = text
            k += 2
        return result

    def _restore_and_validate(parsed, attempt_label):
        """
        Восстановить токены, применить Fact Validation,
        применить Quality Gate. Вернуть (id_to_text, block_reports).
        """
        id_to_text = {}
        block_reports = []
        for i, item in enumerate(orig_items):
            iid = item_ids[i]
            orig_text = item["text"]
            # Наследуем реально применённое решение первой попытки —
            # не пересчитываем через _classify_item заново, иначе
            # hard-frozen блоки (idx<=1) теряют freeze-статус.
            strategy = strategy_map.get(iid, "freeze") if iid in parsed else "freeze"

            if iid not in parsed or strategy == "freeze":
                id_to_text[iid] = orig_text
                block_reports.append({
                    "id": iid, "attempt": attempt_label,
                    "strategy": strategy,
                    "decision": "kept_original",
                    "reason": "frozen_or_missing",
                    "similarity": 1.0,
                })
                continue

            improved = _restore_text(parsed[iid], store)

            # Fact Validation
            fact_ok, fact_reason = _validate_block(orig_text, improved)
            if not fact_ok:
                id_to_text[iid] = orig_text
                block_reports.append({
                    "id": iid, "attempt": attempt_label,
                    "strategy": strategy,
                    "decision": "rejected_facts",
                    "reason": fact_reason,
                    "similarity": _text_similarity(orig_text, improved),
                })
                continue

            # Quality Gate
            qg_ok, sim, qg_reason = _quality_gate(orig_text, improved)
            id_to_text[iid] = improved if qg_ok else None  # None = нужен retry
            block_reports.append({
                "id": iid, "attempt": attempt_label,
                "strategy": strategy,
                "decision": "accepted" if qg_ok else "needs_retry",
                "reason": qg_reason,
                "similarity": sim,
            })

        return id_to_text, block_reports

    # --- Первая попытка ---
    parsed_1 = _parse_ai_response(raw_response)
    id_to_text_1, reports_1 = _restore_and_validate(parsed_1, "attempt_1")

    # Блоки которые нужно переделать (strategy=improve, quality gate не прошли)
    retry_ids = [r["id"] for r in reports_1 if r["decision"] == "needs_retry"]
    all_reports = reports_1

    tokens_total = tokens

    if retry_ids:
        # --- Усиленный промпт для повторной попытки ---
        retry_items_text = []
        for iid in retry_ids:
            idx_r = int(iid) - 1
            orig_t = orig_items[idx_r]["text"] if idx_r < len(orig_items) else ""
            if strategy_map.get(iid) == "freeze":
                # Наследуем freeze-решение первой попытки — не отправлять
                # hard-frozen блок в _protect_text(), иначе обычный текст
                # (например имя) уходит в LLM незащищённым.
                tok = _make_token()
                store[tok] = orig_t
                protected_t = tok
            else:
                protected_t = _protect_text(orig_t, store)
            retry_items_text.append(f"###ITEM_{iid}###\n{protected_t}")

        retry_input = "\n\n".join(retry_items_text)
        n_retry = len(retry_ids)

        retry_system = (
            f"You are a professional resume editor. These {n_retry} resume blocks were sent to you before "
            f"and returned nearly identical to the original — but that is only acceptable if there is "
            f"genuinely nothing to improve. Try again, but only change wording that can be genuinely "
            f"strengthened:\n\n"
            f"- Use a stronger, more precise action verb ONLY if a better one exists for what is already described\n"
            f"- NEVER replace a word with a weaker or more generic synonym (e.g. \"Collaborated with\" → "
            f"\"Worked with\" is FORBIDDEN). If unsure the replacement is stronger, keep the original word.\n"
            f"- Do NOT add new clauses, outcomes, or explanations of impact\n"
            f"- Do NOT add causal or result phrases (\"resulting in\", \"which improved\", \"leading to\", "
            f"\"by leveraging\", \"ensuring\", \"driving\", \"informing\")\n"
            f"- Write in {detected_lang} only\n"
            f"- Copy protected tokens @@@...@@@ EXACTLY as-is\n"
            f"- Return EXACTLY {n_retry} blocks with ###ITEM_NNN### identifiers\n"
            f"- NEVER invent new facts, numbers, companies or technologies\n"
            f"- If a block genuinely has nothing to improve (e.g. it is a name, title, header, or list of "
            f"items), return it completely unchanged — do not pad it to appear different."
        )
        retry_user = (
            f"Rewrite these {n_retry} blocks only where a genuine wording improvement is possible. "
            f"If a block has nothing meaningful to improve, return it unchanged. "
            f"Return with ###ITEM_NNN### identifiers.\n\n"
            f"{retry_input}\n\nOUTPUT ({n_retry} blocks):"
        )

        retry_payload = {
            "model": "openai/gpt-oss-120b",
            "messages": [
                {"role": "system", "content": retry_system},
                {"role": "user", "content": retry_user},
            ],
            "temperature": 0.6,  # выше температура для более творческого ответа
            "max_tokens": 4000,
        }

        resp2 = req_lib.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
            json=retry_payload, timeout=90,
        )

        if resp2.status_code == 429:
            retry_payload["model"] = "openai/gpt-oss-20b"
            resp2 = req_lib.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers={"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"},
                json=retry_payload, timeout=90,
            )

        if resp2.status_code == 200:
            raw_response_2 = resp2.json()["choices"][0]["message"]["content"]
            tokens_total += resp2.json().get("usage", {}).get("total_tokens", 0)
            parsed_2 = _parse_ai_response(raw_response_2)
            # Один вызов — используем один и тот же результат и для отчёта,
            # и для merge, чтобы то, что проверялось, гарантированно совпадало
            # с тем, что попадает в итоговый текст.
            id_to_text_retry, reports_2 = _restore_and_validate(parsed_2, "attempt_2")

            for iid in retry_ids:
                if id_to_text_retry.get(iid) is not None:
                    id_to_text_1[iid] = id_to_text_retry[iid]
                else:
                    # Вторая попытка тоже не прошла — оставляем оригинал
                    idx_r = int(iid) - 1
                    id_to_text_1[iid] = orig_items[idx_r]["text"] if idx_r < len(orig_items) else ""

            all_reports += reports_2

    # --- Финальная сборка ---
    restored_list = []
    for i, item in enumerate(orig_items):
        iid = item_ids[i]
        val = id_to_text_1.get(iid)
        restored_list.append(val if val is not None else item["text"])

    # [DEBUG-LEAK-BULLETS] Диагностика: что лежит в restored_list для
    # буллет-блоков сразу после restore/validate, до сборки improved_text_for_docx.
    try:
        from flask import current_app
        for _i, _val in enumerate(restored_list):
            if '•' in _val:
                current_app.logger.info(
                    "[DEBUG-LEAK-BULLETS] restored_list[%s] item_id=%s repr=%r newline_count=%s",
                    _i, item_ids[_i], _val, _val.count(chr(10)),
                )
    except Exception:
        pass

    # Цикл A: суффикс :TYPE добавляется ТОЛЬКО здесь, на финальной сборке —
    # LLM все ###ITEM_NNN### маркеры (в ai_blocks, в retry, в парсинге
    # ответа модели) видел и продолжает видеть без суффикса. Тип берётся
    # из _extract_structured(); для fallback-пути без docx всегда "plain".
    improved_text_for_docx = (
        "###ITEM_" +
        "\n\n###ITEM_".join(
            f"{item_ids[i]}:{orig_items[i].get('type', 'plain').upper()}###\n{restored_list[i]}"
            for i in range(len(orig_items))
        )
    )
    display_text = "\n".join(restored_list)

    try:
        from flask import current_app
        current_app.logger.info(
            "[DEBUG-TYPE-SUFFIX] improved_text_for_docx repr(first 300): %r",
            improved_text_for_docx[:300],
        )
    except Exception:
        pass

    # --- Отчёт по блокам ---
    quality_report = {
        "total_blocks": len(orig_items),
        "retry_triggered": len(retry_ids),
        "retry_ids": retry_ids,
        "blocks": all_reports,
        "summary": {
            "accepted":        sum(1 for r in all_reports if r["decision"] == "accepted"),
            "kept_original":   sum(1 for r in all_reports if r["decision"] == "kept_original"),
            "rejected_facts":  sum(1 for r in all_reports if r["decision"] == "rejected_facts"),
            "needs_retry":     sum(1 for r in all_reports if r["decision"] == "needs_retry"),
            "avg_similarity":  round(
                sum(r["similarity"] for r in all_reports) / max(len(all_reports), 1), 3
            ),
        },
    }

    return {
        "success": True,
        "improved_resume": improved_text_for_docx,
        "display_text": display_text,
        "original_text": resume_text,
        "detected_language": detected_lang,
        "tokens_used": tokens_total,
        "has_original_docx": original_bytes is not None,
        "quality_report": quality_report,
        "item_ids": item_ids,
    }


def register_missing_routes(app, _extract_text_from_request, _get_current_user):
    from flask import request, jsonify, session, send_file
    import io


    @app.route("/api/admin/improve", methods=["POST"])
    def legacy_admin_improve():
        if "admin" not in session:
            return jsonify({"error": "Unauthorized"}), 401
        try:
            from flask import current_app

            api_key = current_app.config.get("GROQ_API_KEY")

            original_bytes = None
            filename = None
            resume_text_fallback = None

            file = request.files.get("file") or request.files.get("resume")
            if file:
                filename = file.filename
                original_bytes = file.read()
            else:
                data = request.get_json() or {}
                resume_text_fallback = data.get("resume_text", "").strip()

            current_app.logger.info(
                "[DEBUG-LEAK] legacy_admin_improve: has_file=%s filename=%s fallback_preview=%r",
                bool(file),
                filename,
                (resume_text_fallback or "")[:80],
            )

            result = _run_improve_pipeline(original_bytes, filename, resume_text_fallback, api_key)

            if not result.get("success"):
                return jsonify({"success": False, "error": result.get("error")}), result.get("status", 500)

            if original_bytes and filename and filename.lower().endswith('.docx'):
                session["original_docx_token"] = _save_temp_upload(original_bytes)
                session["item_ids"] = result["item_ids"]

            return jsonify(result)

        except Exception as e:
            import traceback
            app.logger.error("legacy_admin_improve failed: %s\n%s", e, traceback.format_exc())
            return jsonify({"success": False, "error": str(e)}), 500

    @app.route("/api/admin/improve/docx", methods=["POST"])
    def legacy_admin_improve_docx():
        if "admin" not in session:
            return jsonify({"error": "Unauthorized"}), 401
        try:
            from docx import Document
            from docx.shared import Pt

            original_file = request.files.get("original_file")
            improved_text = request.form.get("improved_resume") or ""

            if not improved_text:
                data = request.get_json() or {}
                improved_text = data.get("improved_resume", "")

            if not improved_text:
                return jsonify({"success": False, "error": "No text provided"}), 400

            # item_ids: сначала из FormData (надёжно), потом из session (fallback)
            item_ids_raw = request.form.get("item_ids") or ""
            if item_ids_raw:
                import json as _json
                try:
                    item_ids = _json.loads(item_ids_raw)
                except Exception:
                    item_ids = []
            else:
                item_ids = session.get("item_ids", [])

            from flask import current_app
            _dbg_token = session.get("original_docx_token")
            _dbg_token_found = None
            if not original_file and _dbg_token:
                try:
                    _dbg_token_found = _load_temp_upload(_dbg_token) is not None
                except Exception:
                    _dbg_token_found = False
            current_app.logger.info(
                "[DEBUG-LEAK] legacy_admin_improve_docx: has_original_file=%s item_ids_source=%s item_ids=%s "
                "session_token_present=%s session_token_preview=%s temp_upload_found=%s",
                bool(original_file),
                "form" if item_ids_raw else "session",
                item_ids,
                bool(_dbg_token),
                (_dbg_token[:12] if _dbg_token else None),
                _dbg_token_found,
            )

            # Defensive: одноразовое использование session-токена. Очищаем
            # сразу после чтения — даже если этот конкретный запрос его не
            # использует (original_file пришёл напрямую), унаследованные от
            # предыдущего запроса token/item_ids не должны пережить этот
            # вызов и попасть в следующий, не связанный с ним запрос.
            session.pop("original_docx_token", None)
            session.pop("item_ids", None)

            if original_file:
                buf = _apply_improved_text_to_docx(original_file.read(), improved_text, item_ids)
                return send_file(buf, as_attachment=True, download_name="improved_resume.docx",
                    mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            token = _dbg_token
            file_bytes = _load_temp_upload(token)
            if file_bytes:
                buf = _apply_improved_text_to_docx(file_bytes, improved_text, item_ids)
                return send_file(buf, as_attachment=True, download_name="improved_resume.docx",
                    mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

            # Fallback — простой текстовый DOCX
            doc = Document()
            doc.styles["Normal"].font.size = Pt(11)
            clean = re.sub(r"###ITEM_\d+###", "", improved_text)
            for line in clean.split(chr(10)):
                line = line.strip().lstrip("#").replace("**", "").replace("*", "").strip()
                doc.add_paragraph(line if line else "")
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return send_file(buf, as_attachment=True, download_name="improved_resume.docx",
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        except Exception as e:
            return jsonify({"success": False, "error": str(e)}), 500


# ---------------------------------------------------------------------------
# ODT-экспорт (Cycle O1) — только функция генерации, без Flask-роута.
# Роут добавляется отдельным циклом (O2) после сверки с текущей логикой
# списания кредитов в app/__init__.py.
# ---------------------------------------------------------------------------

def _generate_odt(text):
    """
    Сгенерировать .odt из текста improved_resume (с маркерами
    ###ITEM_NNN### или ###ITEM_NNN:TYPE###). ODT не привязан к структуре
    оригинала (в отличие от DOCX-пути, где восстановление идёт по
    item_ids) — маркеры используются здесь только чтобы определить тип
    блока для форматирования, сам ID не нужен.

    Цикл C: раньше (Cycle O1) суффикс :TYPE отбрасывался вместе с самим
    маркером (re.sub схлопывал весь ###ITEM_NNN(:TYPE)?### в перевод
    строки) — тип терялся, все блоки рендерились одинаковым обычным
    абзацем. Теперь тип каждого блока сохраняется (тот же re.split с
    захватывающей группой, что уже применён в _generate_rtf, Цикл B,
    Шаг 2 — переиспользован без изменений) и определяет форматирование
    через нативные именованные стили odfpy:
      - HEADING — жирный текст, чуть увеличенный размер шрифта,
        отступ сверху (аналог \\sb в RTF)
      - TABLE   — жирный текст, БЕЗ дополнительного отступа (пары
        "должность+даты" должны визуально смотреться как единый блок)
      - BULLET  — символ буллета (•) перед текстом + левый отступ
        параграфа
      - PLAIN   — без изменений, как раньше (Standard-стиль)

    В odfpy стили форматирования объявляются заранее как именованные
    объекты style.Style в office:automatic-styles и применяются через
    text:style-name у параграфа целиком — не как inline-атрибуты на
    каждый отдельный ран текста. Здесь текст всегда идёт прямым
    содержимым параграфа (без <text:span>), поэтому style:text-properties
    именованного паражрафного стиля действует как форматирование по
    умолчанию для этого текста — того же эффекта, что и явный span,
    без необходимости в нём.

    RTL (иврит, writing-mode: rl-tb) — как и раньше, определяется по
    диапазону символов \\u0590-\\u05FF и комбинируется с типом блока:
    для каждого типа заведена отдельная RTL-версия стиля (и заголовка,
    и таблицы, и буллета), поскольку у параграфа может быть только один
    именованный стиль одновременно.

    Проверено round-trip тестом (запись -> odf.opendocument.load -> чтение):
    см. tests/test_odt_export.py.

    ВАЖНО (найдено при верификации, не было в исходном наброске):
    класс называется OpenDocumentText, а не OpendocumentText — с опечаткой
    импорт падает немедленно. automaticstyles.addElement() — правильный
    метод для регистрации автоматического стиля абзаца, менять не пришлось.
    """
    import re
    import io
    try:
        from flask import current_app
        current_app.logger.info("[DEBUG-EXPORT-SPACING] %s received text repr: %r", "_generate_odt", text[:300])
    except Exception:
        pass
    from odf.opendocument import OpenDocumentText
    from odf.text import P
    from odf.style import Style, ParagraphProperties, TextProperties

    doc = OpenDocumentText()

    # ------------------------------------------------------------------
    # Именованные стили по типу блока x RTL/LTR (Цикл C). Заведены
    # заранее (не лениво по ходу обхода строк), чтобы каждое имя стиля
    # регистрировалось в automaticstyles ровно один раз — повторная
    # регистрация стиля с тем же именем на каждой строке дала бы
    # дублирующиеся <style:style> с одинаковым style:name, что как
    # минимум избыточно раздувает файл, а с некоторыми ODF-читалками
    # может дать неопределённое поведение (какой из дублей применяется).
    # ------------------------------------------------------------------
    def _make_odt_style(name, is_rtl, bold=False, font_size=None,
                          margin_top=None, margin_left=None):
        para_kwargs = {}
        if is_rtl:
            para_kwargs["writingmode"] = "rl-tb"
        if margin_top is not None:
            para_kwargs["margintop"] = margin_top
        if margin_left is not None:
            para_kwargs["marginleft"] = margin_left

        style = Style(name=name, family="paragraph")
        if para_kwargs:
            style.addElement(ParagraphProperties(**para_kwargs))

        text_kwargs = {}
        if bold:
            text_kwargs["fontweight"] = "bold"
        if font_size:
            text_kwargs["fontsize"] = font_size
        if text_kwargs:
            style.addElement(TextProperties(**text_kwargs))

        doc.automaticstyles.addElement(style)
        return style

    # (block_type, is_rtl) -> Style | None. PLAIN+LTR остаётся None —
    # это Standard-стиль читалки по умолчанию, как и было до Цикла C.
    _odt_styles = {("PLAIN", False): None}
    for is_rtl, suffix in ((False, ""), (True, "RTL")):
        if is_rtl:
            _odt_styles[("PLAIN", True)] = _make_odt_style(
                f"Plain{suffix}", True
            )
        _odt_styles[("HEADING", is_rtl)] = _make_odt_style(
            f"Heading{suffix}", is_rtl, bold=True, font_size="13pt",
            margin_top="0.35cm",
        )
        _odt_styles[("TABLE", is_rtl)] = _make_odt_style(
            f"Table{suffix}", is_rtl, bold=True,
        )
        _odt_styles[("BULLET", is_rtl)] = _make_odt_style(
            f"Bullet{suffix}", is_rtl, margin_left="0.5cm",
        )

    # CRLF-нормализация ДО разбора маркеров — не трогать (см. докстринг
    # выше и коммит 7336c10 для DOCX-пути, тот же баг/то же решение).
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Разбор маркеров с сохранением типа блока — идентичная логика
    # _generate_rtf (Цикл B, Шаг 2), переиспользована без изменений.
    # Markerless-вход (без ###ITEM### вообще) даёт parts из одного
    # элемента — весь текст обрабатывается как единственный PLAIN-блок.
    parts = re.split(r"###ITEM_\d+(?::(\w+))?###", text)
    if len(parts) == 1:
        blocks = [("PLAIN", parts[0])]
    else:
        blocks = []
        for i in range(1, len(parts), 2):
            block_type = (parts[i] or "PLAIN").upper()
            content = parts[i + 1] if i + 1 < len(parts) else ""
            blocks.append((block_type, content))

    for block_type, content in blocks:
        if block_type not in ("HEADING", "TABLE", "BULLET"):
            block_type = "PLAIN"
        content = content.strip("\n")
        for line in content.split("\n"):
            line = line.strip().lstrip("#").replace("**", "").replace("*", "").strip()

            if not line:
                # Пустая строка (межблочный/внутриблочный разделитель) —
                # обычный пустой абзац без стиля, как и в старой версии
                # (Cycle O1: stylename=None для пустых строк).
                doc.text.addElement(P(text=None))
                continue

            has_hebrew = any("\u0590" <= c <= "\u05FF" for c in line)
            display_text = f"\u2022 {line}" if block_type == "BULLET" else line
            style = _odt_styles[(block_type, has_hebrew)]
            doc.text.addElement(P(stylename=style, text=display_text))

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# PDF-экспорт (Cycle P1) — только функция генерации, без Flask-роута.
# Роут добавляется отдельным циклом (P2), фронтенд (замена текущего
# фейкового .pdf-пути в downloadImproved()) — циклом (P3).
# ---------------------------------------------------------------------------

import os as _os_for_font_path

# Alef (SIL OFL 1.1) — static Regular instance, google/fonts repo
# (ofl/alef/Alef-Regular.ttf). Взят вместо Noto Sans Hebrew: в google/fonts
# Noto Sans Hebrew существует ТОЛЬКО как variable font — там нет static/
# подпапки для этого семейства. Регистрация variable-шрифта в reportlab
# технически проходит (reportlab читает default master), но это скрытая
# версийная зависимость того же класса риска, что уже ловили на поведении
# конкретной версии python-docx с merge-ячейками (см.
# tests/test_merged_cells_regression.py) — не хотим повторять паттерн.
# Alef — настоящий static TTF, целиком под иврит, лицензия подтверждена
# по факту скачанного файла: static/fonts/Alef-OFL.txt (SIL OFL 1.1,
# явно разрешает "bundled, embedded, redistributed and/or sold with any
# software").
_PDF_FONT_NAME = "Alef"
_PDF_FONT_PATH = _os_for_font_path.path.join(
    _os_for_font_path.path.dirname(_os_for_font_path.path.dirname(_os_for_font_path.path.abspath(__file__))),
    "static", "fonts", "Alef-Regular.ttf",
)

# Fira Sans (SIL OFL 1.1) — static Regular instance, тот же google/fonts репо
# (ofl/firasans/FiraSans-Regular.ttf), тот же принцип выбора: настоящий
# static TTF, не variable-инстанс.
# Причина добавления: Alef покрывает иврит и, неожиданно, расширенную
# латиницу (é/ñ/ł/č/ü и т.п. — проверено через fontTools.getBestCmap на
# всех cmap-подтаблицах), но НЕ содержит ни одного кириллического глифа —
# то есть PDF для уже заявленных на сайте русского/украинского резюме молча
# рисовал пустые квадраты вместо букв. Fira Sans проверена так же (все
# cmap-подтаблицы): полное покрytие кириллицы + той же расширенной
# латиницы. Используется как LTR-шрифт по умолчанию для всех НЕ-ивритских
# строк; Alef остаётся только для строк с ивритскими символами (та же
# эвристика \\u0590-\\u05FF, что уже была).
_PDF_FONT_NAME_LATIN = "FiraSans"
_PDF_FONT_PATH_LATIN = _os_for_font_path.path.join(
    _os_for_font_path.path.dirname(_os_for_font_path.path.dirname(_os_for_font_path.path.abspath(__file__))),
    "static", "fonts", "FiraSans-Regular.ttf",
)

# Noto Sans Arabic (SIL OFL 1.1) — google/fonts repo (ofl/notosansarabic/
# NotoSansArabic[wdth,wght].ttf). ОТКЛОНЕНИЕ от принципа "всегда static TTF"
# (см. комментарий у Alef выше) — проверено эмпирически и задокументировано
# намеренно, а не по недосмотру:
#   - У Noto Sans Arabic/CJK НЕТ static-инстансов в формате, который вообще
#     принимает reportlab. В noto-cjk репозитории (аналог для арабского не
#     проверялся отдельно, но тот же паттерн подтверждён на китайском —
#     см. отчёт по шагу 1 для CJK) static-сборки существуют только как
#     OTF/CFF (PostScript outlines) — reportlab.pdfbase.ttfonts.TTFont их
#     физически не грузит: TTFError "postscript outlines are not supported".
#     Проверено напрямую попыткой регистрации, не предположение.
#   - Единственный формат, который reportlab принимает (glyf/TrueType) —
#     только variable font. Риск, из-за которого Alef/FiraSans выбирались
#     static (непредсказуемый default instance — см. ниже пример с
#     китайским, где default оказался Thin=100), здесь ПРОВЕРЕН и не
#     воспроизвёлся: у NotoSansArabic[wdth,wght].ttf default инстанс оси
#     wght = 400 (Regular), что подтверждено через fontTools (f['fvar'].axes)
#     и визуальной проверкой отрендеренного PDF. Файл сохранён под именем
#     NotoSansArabic-VF.ttf (суффикс VF — variable font) умышленно, чтобы
#     не выдавать его за static-инстанс при будущих ревью кода.
#   - Вес ~825 КБ — некритично (для сравнения FiraSans ~450 КБ).
_PDF_FONT_NAME_ARABIC = "NotoArabic"
_PDF_FONT_PATH_ARABIC = _os_for_font_path.path.join(
    _os_for_font_path.path.dirname(_os_for_font_path.path.dirname(_os_for_font_path.path.abspath(__file__))),
    "static", "fonts", "NotoSansArabic-VF.ttf",
)

# WenQuanYi Micro Hei (Apache License 2.0 / GPLv3-with-font-embedding-
# exception, dual license — см. static/fonts/WenQuanYiMicroHei-Apache2-
# LICENSE.txt и WenQuanYiMicroHei-GPLv3-LICENSE.txt, оба файла — реально
# скачанные тексты лицензий из official source repo, не пересказ по памяти).
# ОТКЛОНЕНИЕ от принципа "всегда static TTF", но в ДРУГУЮ сторону, чем у
# NotoSansArabic-VF выше: здесь как раз НАЙДЕН настоящий static-инстанс —
# проверено (Cycle CN — диагностика и research), что готовые CJK static
# сборки Noto (OTF, оба найденных файла 8МБ/16МБ) физически не грузятся
# reportlab.pdfbase.ttfonts.TTFont: TTFError "postscript outlines are not
# supported" (CFF-контуры, не glyf). Единственный найденный glyf-вариант
# от Noto — variable font с default wght=100 (Thin), то есть "ложный успех"
# по аналогии с риском, уже пойманным у Alef/FiraSans (см. их комментарии
# выше) — здесь default НЕ подошёл бы, в отличие от NotoSansArabic-VF, где
# default оказался верным 400/Regular.
# WenQuanYi Micro Hei снимает обе проблемы разом: это подлинный static TTF
# (не .ttc — оригинальный файл распространяется как TrueType Collection
# из двух начертаний, "Micro Hei" и "Micro Hei Mono"; здесь извлечён и
# сохранён ТОЛЬКО первый subfont через fontTools — TTCollection.fonts[0].
# save(...) — как отдельный самостоятельный .ttf), контуры glyf, вес Regular
# подтверждён из name table (nameID=2 "Regular") и визуальной проверкой
# (штрих сопоставимой толщины с латинской строкой на том же рендере — не
# тонкий/hairline, как было бы у ложного default-инстанса variable font).
# Покрытие: 20 932 CJK-глифа в диапазоне \\u4E00-\\u9FFF (полное для целей
# резюме). Вес файла — 4.4 МБ, для сравнения: NotoSansArabic-VF ~825 КБ,
# FiraSans ~450 КБ — WenQuanYi заметно тяжелее остальных трёх шрифтов
# проекта вместе взятых, но это ожидаемо для полного CJK-шрифта (тысячи
# уникальных глифов против десятков latin/cyrillic/arabic) и остаётся
# приемлемым для бесплатного тарифа Render.
_PDF_FONT_NAME_CJK = "WenQuanYiMicroHei"
_PDF_FONT_PATH_CJK = _os_for_font_path.path.join(
    _os_for_font_path.path.dirname(_os_for_font_path.path.dirname(_os_for_font_path.path.abspath(__file__))),
    "static", "fonts", "WenQuanYiMicroHei-Regular.ttf",
)

_pdf_font_registered = False


def _ensure_pdf_font_registered():
    """Зарегистрировать все PDF-шрифты (Alef — иврит, FiraSans — всё
    остальное включая кириллицу, NotoArabic — арабский, WenQuanYiMicroHei —
    китайский) в reportlab один раз за процесс."""
    global _pdf_font_registered
    if _pdf_font_registered:
        return
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    pdfmetrics.registerFont(TTFont(_PDF_FONT_NAME, _PDF_FONT_PATH))
    pdfmetrics.registerFont(TTFont(_PDF_FONT_NAME_LATIN, _PDF_FONT_PATH_LATIN))
    pdfmetrics.registerFont(TTFont(_PDF_FONT_NAME_ARABIC, _PDF_FONT_PATH_ARABIC))
    pdfmetrics.registerFont(TTFont(_PDF_FONT_NAME_CJK, _PDF_FONT_PATH_CJK))
    _pdf_font_registered = True


def _generate_pdf(text):
    """
    Сгенерировать .pdf из текста improved_resume (с маркерами ###ITEM_NNN###
    или ###ITEM_NNN:TYPE###) — PDF не привязан к структуре оригинала,
    восстановление по item_ids здесь не нужно, как и в _generate_odt.

    Цикл D2: раньше (Cycle P1) суффикс :TYPE отбрасывался вместе с самим
    маркером одним re.sub (r"\\n*###ITEM_\\d+(?::\\w+)?###\\n*" -> "\\n") —
    тип терялся, все блоки рисовались одинаковым обычным текстом. Теперь
    маркер разбирается через re.split с сохранением типа (тот же паттерн,
    что уже применён в _generate_rtf/_generate_odt, Циклы B/C — переиспользован
    без изменений) и определяет форматирование:
      - HEADING — увеличенный font_size (+2pt, тот же шаг что и в ODT —
        13pt вместо базовых 11pt) + дополнительный отступ сверху перед
        первой строкой блока (аналог \\sb в RTF / margintop в ODT)
      - TABLE   — увеличенный font_size (+2pt), БЕЗ доп. отступа (пары
        "должность+даты" после Цикла B, Шаг 1 идут подряд и должны
        визуально смотреться как единый связанный блок — тот же принцип,
        что уже применён в RTF/ODT)
      - BULLET  — литеральный символ буллета "• " перед строкой, обычный
        font_size, без изменения шрифта
      - PLAIN   — без изменений, как было раньше

    РЕШЕНИЕ по bold (принято перед началом цикла, не блокер для этого
    цикла): в проекте зарегистрированы только Regular-варианты всех
    четырёх PDF-шрифтов (Alef/FiraSans/NotoArabic/WenQuanYiMicroHei) —
    ни одного Bold TTF-файла нет. В ReportLab с кастомным TTF через
    pdfmetrics.registerFont(TTFont(...)) жирность — это не атрибут рана,
    а отдельный шрифтовой файл; без него bold=True работает только для
    встроенных core-шрифтов (Helvetica-Bold и т.п.), не для наших TTF.
    Вместо добавления Bold-файлов в этом цикле выбран более дешёвый путь —
    увеличенный размер шрифта без реального bold. Из-за этого HEADING/TABLE
    в PDF визуально слабее отделены от PLAIN, чем в RTF/ODT (там настоящий
    bold) — это осознанное отклонение от полного визуального паритета
    между тремя экспортами, не оставшийся баг. Полный паритет (поиск/
    добавление настоящих Bold-шрифтов по той же дисциплине проверки
    static-vs-variable, что уже применена для Regular-вариантов) можно
    рассмотреть отдельным циклом позже, если понадобится.

    Перенос строк (_wrap) измеряет ширину через pdfmetrics.stringWidth с
    font_size, ПЕРЕДАННЫМ как параметр конкретного блока (11pt PLAIN/
    BULLET, 13pt HEADING/TABLE) — не захардкоженной переменной из
    замыкания, иначе HEADING/TABLE переносились бы по ширине, рассчитанной
    для меньшего шрифта, и текст визуально вылезал бы за правое поле
    страницы на PLAIN-подобной ширине переноса. Подтверждено эмпирически
    (изолированный скрипт со встроенным Helvetica, т.к. кастомные TTF
    недоступны вне рабочего окружения): при одинаковой max_width перенос
    для 13pt даёт другое разбиение по словам, чем для 11pt — параметр
    реально учитывается, не игнорируется.

    RTL: строки, содержащие иврит (диапазон \\u0590-\\u05FF, тот же что и
    везде в проекте), прогоняются через bidi.algorithm.get_display()
    (логический порядок символов -> визуальный порядок для LTR-рендеринга)
    и рисуются выравниванием по правому краю (drawRightString), шрифтом
    Alef. Строки с арабским (диапазон \\u0600-\\u06FF) — тем же
    drawRightString, шрифтом NotoArabic, НО перед get_display() дополнительно
    прогоняются через arabic_reshaper.reshape(). Это не опечатка и не лишний
    шаг: иврит и арабский оба RTL, но только у арабского буквы физически
    меняют начертание (contextual shaping: изолированная/начальная/
    срединная/конечная форма + лигатуры типа lam-alef) в зависимости от
    соседних букв. Без reshape() reportlab рисует каждую букву в
    изолированной форме — технически не ошибка на уровне кода (PDF
    валиден), но визуально неверный, нечитаемый для носителя языка текст.
    Проверено визуально (рендер в PNG) на обоих вариантах — без reshape
    буквы стоят раздельно, с reshape — слитно, как положено. У иврита
    такой проблемы нет: ивритские буквы не меняют форму от соседей, только
    порядок написания (bidi без shaping — стандартное и достаточное
    решение для иврита, но недостаточное для арабского).
    Китайский: строки, содержащие символы диапазона \\u4E00-\\u9FFF,
    рисуются обычным drawString слева (LTR — не RTL, bidi/reshape не
    применяются, в отличие от иврита/арабского выше), шрифтом
    WenQuanYiMicroHei. См. комментарий у _PDF_FONT_NAME_CJK — единственный
    из четырёх шрифтов проекта, для которого удалось найти подлинный
    static TTF без необходимости идти на variable-font компромисс (как
    пришлось для NotoSansArabic-VF).

    Остальные строки (латиница, кириллица) — обычным drawString слева,
    шрифтом FiraSans (см. комментарий у _PDF_FONT_NAME_LATIN — Alef
    кириллицу не содержит).

    Перенос строк: greedy word-wrap по словам, измерение через
    pdfmetrics.stringWidth на логическом (не bidi-переставленном) тексте,
    ширины символов берутся из шрифта КОНКРЕТНОЙ строки (Alef или
    FiraSans — они разные, нельзя мерить одним шрифтом текст, который
    будет нарисован другим) — перестановка применяется ПОСЛЕ переноса,
    к каждой уже готовой под-строке отдельно, иначе разбиение по словам
    съезжает относительно визуального порядка. Не идеальный перенос (не
    бьёт слово посередине, если оно само шире доступной ширины) — этого
    достаточно для резюме.

    Пагинация: при достижении нижнего поля — showPage() + повторная
    setFont() (шрифт не сохраняется между страницами в reportlab).

    ВАЖНО (найдено при верификации, не было в исходном наброске):
    PyPDF2.extract_text() не гарантирует читаемый логический порядок для
    ивритского/bidi-текста даже из корректно сгенерированного PDF — это
    ограничение формата/библиотеки, не специфичное для этой функции.
    Подробности и что реально проверено — см. tests/test_pdf_export.py
    и отчёт цикла P1.
    """
    import re
    import io
    try:
        from flask import current_app
        current_app.logger.info("[DEBUG-EXPORT-SPACING] %s received text repr: %r", "_generate_pdf", text[:300])
    except Exception:
        pass
    from reportlab.pdfgen import canvas as pdf_canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from bidi.algorithm import get_display
    import arabic_reshaper

    _ensure_pdf_font_registered()

    font_hebrew = _PDF_FONT_NAME
    font_latin = _PDF_FONT_NAME_LATIN
    font_arabic = _PDF_FONT_NAME_ARABIC
    font_cjk = _PDF_FONT_NAME_CJK
    font_size = 11
    # HEADING/TABLE — увеличенный размер вместо недоступного настоящего
    # bold (см. докстринг функции, раздел "РЕШЕНИЕ по bold"). Тот же шаг
    # (+2pt), что уже применён для этих же типов блоков в _generate_odt.
    heading_table_font_size = font_size + 2
    line_height = 15
    # Доп. отступ перед первой строкой HEADING-блока — аналог \sb в RTF /
    # margintop в ODT. В этой построчной модели без вложенности "блок ->
    # строки" отступ реализован как довесок к обычному шагу y на первой
    # физической строке блока (Вариант А из диагностики Цикла D1), а не
    # отдельный "пустой шаг" — тот же принцип, что и в RTF (\sb240 — это
    # тоже довесок к уже существующему шагу параграфа, не отдельный проход).
    heading_margin_top = 10
    margin = 50

    page_w, page_h = A4
    max_width = page_w - 2 * margin

    buf = io.BytesIO()
    c = pdf_canvas.Canvas(buf, pagesize=A4)
    c.setFont(font_latin, font_size)
    y = page_h - margin

    def _wrap(line, font_name, size):
        # size передаётся параметром конкретного блока (11pt PLAIN/BULLET,
        # 13pt HEADING/TABLE) — НЕ читается из переменной font_size
        # замыкания, иначе HEADING/TABLE (реально нарисованные большим
        # шрифтом) переносились бы по ширине, посчитанной для меньшего
        # шрифта, и текст вылезал бы за правое поле страницы. Подтверждено
        # эмпирически отдельным скриптом — см. отчёт цикла.
        words = [w for w in line.split(" ") if w != ""] or [""]
        wrapped = []
        current = words[0]
        for w in words[1:]:
            candidate = current + " " + w
            if pdfmetrics.stringWidth(candidate, font_name, size) <= max_width:
                current = candidate
            else:
                wrapped.append(current)
                current = w
        wrapped.append(current)
        return wrapped

    # Цикл D2: разбор маркера с сохранением типа блока — тот же паттерн,
    # что уже применён в _generate_rtf/_generate_odt (Циклы B/C, code
    # переиспользован без изменений). ДО этого исправления (Cycle P1)
    # один re.sub схлопывал ###ITEM_NNN(:TYPE)?### целиком, включая
    # суффикс :TYPE — тип терялся и HEADING/TABLE/BULLET рисовались
    # визуально одинаково с PLAIN. re.split с одной захватывающей группой
    # даёт [текст_до_первого_маркера, type1, content1, type2, content2, ...];
    # markerless-вход (без ###ITEM### вообще, например прямой вызов из
    # теста) даёт parts из одного элемента — весь текст обрабатывается
    # как единственный PLAIN-блок, в точности воспроизводя поведение ДО
    # этого цикла для такого входа.
    # CRLF-нормализация ДО разбора маркера — не трогать (см. докстринг
    # выше и коммит 7336c10 для DOCX-пути, тот же баг/то же решение).
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    parts = re.split(r"###ITEM_\d+(?::(\w+))?###", text)
    if len(parts) == 1:
        blocks = [("PLAIN", parts[0])]
    else:
        blocks = []
        for i in range(1, len(parts), 2):
            block_type = (parts[i] or "PLAIN").upper()
            content = parts[i + 1] if i + 1 < len(parts) else ""
            blocks.append((block_type, content))

    for block_type, content in blocks:
        if block_type not in ("HEADING", "TABLE", "BULLET"):
            block_type = "PLAIN"
        content = content.strip("\n")
        block_font_size = (
            heading_table_font_size if block_type in ("HEADING", "TABLE") else font_size
        )
        # Флаг "первая физическая строка блока" — только на ней (и только
        # для HEADING) применяется доп. отступ сверху. content.strip("\n")
        # выше уже убрал ведущие/замыкающие пустые строки блока, поэтому
        # первым элементом content.split("\n") практически всегда будет
        # непустая строка — это и есть "первая строка блока" в терминах
        # диагностики Цикла D1.
        is_first_line_of_block = True

        for raw_line in content.split("\n"):
            line = raw_line.strip().lstrip("#").replace("**", "").replace("*", "").strip()

            if not line:
                if y < margin + line_height:
                    c.showPage()
                    c.setFont(font_latin, font_size)
                    y = page_h - margin
                y -= line_height
                is_first_line_of_block = False
                continue

            # BULLET — литеральный символ буллета перед строкой, тот же
            # видимый эффект, что "\bullet " в RTF и "\u2022 " в ODT.
            display_line = f"\u2022 {line}" if block_type == "BULLET" else line

            has_hebrew = any("\u0590" <= ch <= "\u05FF" for ch in display_line)
            has_arabic = any("\u0600" <= ch <= "\u06FF" for ch in display_line)
            has_cjk = any("\u4E00" <= ch <= "\u9FFF" for ch in display_line)
            if has_hebrew:
                line_font = font_hebrew
            elif has_arabic:
                line_font = font_arabic
            elif has_cjk:
                line_font = font_cjk
            else:
                line_font = font_latin

            extra_top = heading_margin_top if (block_type == "HEADING" and is_first_line_of_block) else 0
            if extra_top:
                # Отступ должен визуально отделять заголовок ОТ ПРЕДЫДУЩЕГО
                # текста (воздух НАД заголовком), а не раздувать промежуток
                # ПОСЛЕ него. Раз y уменьшается на "step" ПОСЛЕ отрисовки
                # текущей строки (и, значит, влияет на позицию следующей),
                # довесок нужно применить здесь — ДО отрисовки первой
                # строки блока, — а не примешивать его в step вместе с
                # обычным line_height (иначе он сдвигает не ту строку).
                if y < margin + extra_top:
                    c.showPage()
                    c.setFont(font_latin, font_size)
                    y = page_h - margin
                y -= extra_top

            for sub_idx, sub_line in enumerate(_wrap(display_line, line_font, block_font_size)):
                # Шаг всегда line_height — без extra_top: тот довесок уже
                # учтён выше, до входа в этот цикл, и относится к отступу
                # НАД заголовком, а не к шагу между его собственными
                # под-строками или к следующей за ним строке.
                step = line_height
                if y < margin + step:
                    c.showPage()
                    y = page_h - margin
                # Явно перед каждой отрисовкой — предыдущая строка могла
                # использовать другой шрифт (Alef/FiraSans/NotoArabic чередуются
                # по тексту) и/или другой размер (PLAIN/BULLET 11pt против
                # HEADING/TABLE 13pt), reportlab не хранит "текущий" шрифт
                # между строками надёжно для наших целей, поэтому не
                # полагаемся на состояние.
                c.setFont(line_font, block_font_size)
                if has_hebrew:
                    c.drawRightString(page_w - margin, y, get_display(sub_line))
                elif has_arabic:
                    # Арабский требует shaping ДО bidi-переупорядочивания —
                    # см. подробное объяснение в docstring функции выше.
                    shaped = arabic_reshaper.reshape(sub_line)
                    c.drawRightString(page_w - margin, y, get_display(shaped))
                else:
                    c.drawString(margin, y, sub_line)
                y -= step

            is_first_line_of_block = False

    # Явно финализировать последнюю страницу перед save(). Без этого:
    # с кастомным TTF-шрифтом (Alef) страница, на которой ни разу не был
    # вызван drawString/drawRightString (например, полностью пустой/
    # пробельный текст на входе — только пустые строки, только y -=
    # line_height без единого реального рисования), молча пропадает —
    # PdfReader(...).pages даёт 0 страниц вместо ожидаемой минимум одной.
    # Со встроенным Helvetica та же ситуация давала 1 страницу нормально;
    # баг воспроизведён изолированно и подтверждён с обоими вариантами до
    # применения фикса — см. tests/test_pdf_export.py::test_09 и отчёт
    # цикла P1. showPage() здесь безопасен и не создаёт лишнюю пустую
    # страницу в конце: любой showPage() внутри цикла выше всегда сразу
    # продолжается рисованием на новой странице (полноценной строкой или
    # хотя бы сдвигом y для пустой строки), то есть никогда не остаётся
    # "висящим" последним вызовом цикла.
    c.showPage()
    c.save()
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# RTF-экспорт (Cycle R1) — только функция генерации, без Flask-роута.
# Роут — отдельным циклом (R2), фронтенд — циклом (R3).
# ---------------------------------------------------------------------------

def _generate_rtf(text):
    """
    Сгенерировать .rtf из текста improved_resume (с маркерами
    ###ITEM_NNN### или ###ITEM_NNN:TYPE###).

    Цикл B, Шаг 2: раньше суффикс :TYPE отбрасывался вместе с самим
    маркером (regex просто схлопывал весь ###ITEM_NNN(:TYPE)?### в
    ничего) — тип терялся, и HEADING/BULLET/TABLE/PLAIN рендерились
    визуально одинаково. Теперь тип каждого блока сохраняется и
    определяет RTF-форматирование:
      - HEADING — жирный (\\b...\\b0), отступ сверху (\\sb/\\sa)
      - TABLE   — жирный (мини-заголовок), БЕЗ доп. отступа — после
        исправления порядка элементов (Цикл B, Шаг 1) пары
        "должность+даты" идут подряд и должны визуально смотреться
        как единый связанный блок
      - BULLET  — символ буллета (\\bullet) + небольшой левый отступ
      - PLAIN   — без изменений, как было раньше

    В отличие от PDF, направление RTL/LTR в RTF задаётся управляющими
    словами (\\rtlch/\\ltrch), а не физической перестановкой символов —
    python-bidi здесь НЕ используется и не нужен: текст остаётся в
    исходном логическом порядке символов, как он есть в строке.

    Проверено round-trip тестом (запись -> striprtf.rtf_to_text() ->
    сравнение с очищенным исходником): см. tests/test_rtf_export.py.

    ВАЖНО (эмпирическая проверка при разработке Шага 2, для честности
    отчёта): символьное форматирование (\\b/\\b0) и paragraph-properties
    (\\sb/\\sa/\\li) НЕ появляются в тексте, извлекаемом striprtf —
    подтверждено отдельным скриптом до применения фикса к файлу. Символ
    буллета (\\bullet) — единственное добавление в этом шаге, которое
    ДЕЙСТВИТЕЛЬНО меняет извлечённый текст: добавляет видимый префикс
    "• " перед строкой BULLET-блока, которого не было в тексте раньше
    (буллеты в оригинальном .docx рендерятся Word'ом через numPr/
    нумерацию списка, а не как символ в самом тексте параграфа — см.
    _classify_para_type). Это ОСОЗНАННОЕ и ОЖИДАЕМОЕ изменение
    извлекаемого содержимого — это и есть цель Шага 2 ("реальный
    символ буллета"), не побочный эффект. tests/test_rtf_export.py
    обновлён под это явно, см. отчёт цикла.

    ВАЖНО (унаследовано из Цикла R1): в отличие от PDF, где RTL-текст
    при обратном извлечении предсказуемо страдал на смешанных RTL+LTR
    строках (см. отчёт цикла P1), здесь round-trip остаётся точным даже
    на смешанном иврит+цифры+email контенте — потому что RTF не
    переставляет символы физически (в отличие от bidi.get_display(),
    применяемого в _generate_pdf), направление — это только метаданные
    отображения, которые striprtf на чтении игнорирует и просто отдаёт
    исходную последовательность символов.
    """
    import re
    try:
        from flask import current_app
        current_app.logger.info("[DEBUG-EXPORT-SPACING] %s received text repr: %r", "_generate_rtf", text[:300])
    except Exception:
        pass

    def _escape_rtf(s):
        # Экранировать RTF-спецсимволы, затем не-ASCII через \uN? escape.
        out = []
        for ch in s:
            if ch == '\\':
                out.append('\\\\')
            elif ch == '{':
                out.append('\\{')
            elif ch == '}':
                out.append('\\}')
            elif ord(ch) > 127:
                code = ord(ch)
                if code > 32767:
                    code -= 65536
                out.append(f'\\u{code}?')
            else:
                out.append(ch)
        return ''.join(out)

    def _rtf_type_wrap(escaped_text, block_type):
        """
        Обернуть уже экранированный текст строки в character-level RTF
        control words по типу блока. \\b/\\b0 (жирный) не влияет на
        извлекаемый striprtf текст — проверено эмпирически. \\bullet —
        единственный элемент здесь, который добавляет видимый символ в
        извлечённый текст (см. докстринг функции).
        """
        if block_type in ("HEADING", "TABLE"):
            return f"\\b {escaped_text}\\b0 "
        if block_type == "BULLET":
            # Один пробел сразу после \bullet — control-word delimiter,
            # поглощается парсером RTF и не становится видимым символом.
            # Второй пробел — уже буквальный, он и даёт "• text" при
            # извлечении (проверено эмпирически striprtf).
            return f"\\bullet  {escaped_text}"
        return escaped_text

    def _rtf_para_props(block_type):
        """
        Paragraph-level control words (отступ/интервал) по типу блока.
        Не влияют на извлекаемый текст (чистые control words без
        видимого содержимого) — проверено эмпирически.
        """
        if block_type == "HEADING":
            # Пространство сверху и немного снизу — визуально отделяет
            # заголовок секции от предыдущего блока.
            return "\\sb240\\sa60 "
        if block_type == "BULLET":
            # Небольшой левый отступ под буллет-маркер.
            return "\\li360 "
        # TABLE — намеренно БЕЗ отступа (см. докстринг функции: пары
        # "должность+даты" после Шага 1 идут подряд и должны выглядеть
        # как один связанный блок). PLAIN — без изменений, как раньше.
        return ""

    # CRLF-нормализация ДО разбора маркеров: клиент пересылает improved_resume
    # через multipart/form-data (FormData), которое браузер по спецификации HTML5
    # нормализует \n -> \r\n. Без этой нормализации маркерный regex ниже (матчащий
    # только литеральные \n) не видит \r и оставляет мусорные \r\n-разделители между
    # блоками нетронутыми — тот же баг и то же решение, что уже применены в
    # _apply_improved_text_to_docx для DOCX-пути (см. коммит 7336c10).
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Разбор маркеров с сохранением типа блока (Цикл B, Шаг 2). Одна
    # захватывающая группа — сам :TYPE (id блока здесь не нужен, как и
    # в старой версии, где он тоже отбрасывался). re.split с одной
    # группой даёт [текст_до_первого_маркера, type1, content1, type2,
    # content2, ...]; если маркеров нет вовсе (markerless-вызов,
    # например напрямую из теста) — parts состоит из одного элемента,
    # и весь текст обрабатывается как единственный PLAIN-блок, что
    # в точности воспроизводит поведение ДО Шага 2 для такого входа.
    parts = re.split(r"###ITEM_\d+(?::(\w+))?###", text)
    if len(parts) == 1:
        blocks = [("PLAIN", parts[0])]
    else:
        blocks = []
        for i in range(1, len(parts), 2):
            block_type = (parts[i] or "PLAIN").upper()
            content = parts[i + 1] if i + 1 < len(parts) else ""
            blocks.append((block_type, content))

    body = []
    for block_type, content in blocks:
        content = content.strip("\n")
        para_props = _rtf_para_props(block_type)
        for line in content.split("\n"):
            line = line.strip().lstrip('#').replace('**', '').replace('*', '').strip()
            has_hebrew = any('\u0590' <= c <= '\u05FF' for c in line)
            escaped = _escape_rtf(line)
            # Пустую строку (межблочный/внутриблочный разделитель) не
            # оборачиваем в форматирование типа — визуально это ничего
            # не меняет, но так чище и не плодит бессмысленные \b\b0
            # вокруг пустоты.
            wrapped = _rtf_type_wrap(escaped, block_type) if line else escaped
            if has_hebrew:
                body.append(f"\\rtlch\\rtlpar\\qr {para_props}{wrapped}\\par")
            else:
                body.append(f"\\ltrch\\ltrpar\\ql {para_props}{wrapped}\\par")

    rtf = (
        r"{\rtf1\ansi\ansicpg1252\uc1\deff0"
        r"{\fonttbl{\f0\fswiss\fcharset0 Arial;}}"
        r"\f0\fs24 "
        + "\n".join(body)
        + "}"
    )
    return rtf.encode('utf-8')
