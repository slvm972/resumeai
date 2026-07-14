#!/usr/bin/env python3
"""
batch_runner.py — внешний скрипт для регрессионного тестирования ResumeAI.

Не изменяет и не импортирует код проекта. Работает исключительно через
HTTP API запущенного сервера (python run.py на localhost:5000).

Использование:
    python batch_runner.py --input ./test_resumes --output ./batch_results --base-url http://127.0.0.1:5000

Для каждого .docx из input-папки выполняет:
    1. POST /api/login          (admin-логин, один раз на весь batch)
    2. POST /api/admin/analyze  (анализ)
    3. POST /api/admin/improve  (улучшение)
    4. POST /api/admin/improve/docx  (получение итогового DOCX)

Сохраняет в output/<basename>/:
    original.docx           — копия исходного файла
    improved.docx            — итоговый файл после improve/docx
    improve_response.json    — полный JSON-ответ /api/admin/improve
    analyze_response.json    — полный JSON-ответ /api/admin/analyze
    quality_report.json      — quality_report из improve_response (если есть)
    log.txt                  — лог выполнения для этого файла

А также в корне output/:
    batch_summary.json       — сводка по всем файлам (статусы, ошибки, тайминги)
    batch_summary.txt        — человекочитаемая версия сводки
"""

import argparse
import json
import os
import sys
import time
import traceback
from pathlib import Path

try:
    import requests
except ImportError:
    print("ERROR: пакет 'requests' не установлен. Установите: pip install requests")
    sys.exit(1)


def log(logfile, message):
    """Записать строку в лог файла и вывести в консоль."""
    line = f"[{time.strftime('%H:%M:%S')}] {message}"
    print(line)
    if logfile:
        logfile.write(line + "\n")
        logfile.flush()


def login(session, base_url, username, password, logfile):
    """Выполнить admin-логин один раз на весь batch."""
    log(logfile, f"Логин администратора ({username})...")
    resp = session.post(
        f"{base_url}/api/login",
        json={"username": username, "password": password},
        timeout=30,
    )
    if resp.status_code != 200:
        raise RuntimeError(f"Login failed: HTTP {resp.status_code} — {resp.text[:300]}")
    body = resp.json()
    if not body.get("success"):
        raise RuntimeError(f"Login failed: {body}")
    log(logfile, "Логин выполнен успешно.")
    return body


def run_single_resume(session, base_url, docx_path, out_dir, logfile):
    """
    Прогнать один DOCX через полный цикл analyze -> improve -> improve/docx.
    Возвращает dict со статусом и метриками для сводки.
    """
    result = {
        "file": docx_path.name,
        "status": "unknown",
        "steps": {},
        "error": None,
        "timings_sec": {},
    }

    out_dir.mkdir(parents=True, exist_ok=True)

    # Копия оригинала для удобства сравнения
    original_copy = out_dir / "original.docx"
    original_copy.write_bytes(docx_path.read_bytes())

    try:
        # -------------------------------------------------------------
        # Шаг 1: /api/admin/analyze
        # -------------------------------------------------------------
        log(logfile, f"[{docx_path.name}] Шаг 1/3: analyze...")
        t0 = time.time()
        with open(docx_path, "rb") as f:
            resp = session.post(
                f"{base_url}/api/admin/analyze",
                files={"file": (docx_path.name, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                timeout=120,
            )
        result["timings_sec"]["analyze"] = round(time.time() - t0, 2)

        if resp.status_code != 200:
            result["status"] = "FAILED_ANALYZE"
            result["error"] = f"HTTP {resp.status_code}: {resp.text[:500]}"
            log(logfile, f"[{docx_path.name}] analyze FAILED: {result['error']}")
            return result

        analyze_body = resp.json()
        (out_dir / "analyze_response.json").write_text(
            json.dumps(analyze_body, indent=2, ensure_ascii=False), encoding="utf-8"
        )
        result["steps"]["analyze"] = "OK"
        log(logfile, f"[{docx_path.name}] analyze OK ({result['timings_sec']['analyze']}s)")

        # -------------------------------------------------------------
        # Шаг 2: /api/admin/improve
        # -------------------------------------------------------------
        log(logfile, f"[{docx_path.name}] Шаг 2/3: improve...")
        t0 = time.time()
        with open(docx_path, "rb") as f:
            resp = session.post(
                f"{base_url}/api/admin/improve",
                files={"file": (docx_path.name, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                timeout=180,
            )
        result["timings_sec"]["improve"] = round(time.time() - t0, 2)

        if resp.status_code != 200:
            result["status"] = "FAILED_IMPROVE"
            result["error"] = f"HTTP {resp.status_code}: {resp.text[:500]}"
            log(logfile, f"[{docx_path.name}] improve FAILED: {result['error']}")
            return result

        improve_body = resp.json()
        (out_dir / "improve_response.json").write_text(
            json.dumps(improve_body, indent=2, ensure_ascii=False), encoding="utf-8"
        )
        result["steps"]["improve"] = "OK"
        result["tokens_used"] = improve_body.get("tokens_used")
        result["detected_language"] = improve_body.get("detected_language")

        quality_report = improve_body.get("quality_report")
        if quality_report:
            (out_dir / "quality_report.json").write_text(
                json.dumps(quality_report, indent=2, ensure_ascii=False), encoding="utf-8"
            )
            result["quality_summary"] = quality_report.get("summary")

        log(logfile, f"[{docx_path.name}] improve OK ({result['timings_sec']['improve']}s), "
                      f"tokens={result.get('tokens_used')}, lang={result.get('detected_language')}")
        if quality_report:
            log(logfile, f"[{docx_path.name}] quality_report.summary = {quality_report.get('summary')}")

        # -------------------------------------------------------------
        # Шаг 3: /api/admin/improve/docx
        # -------------------------------------------------------------
        log(logfile, f"[{docx_path.name}] Шаг 3/3: improve/docx...")
        t0 = time.time()

        improved_resume_text = improve_body.get("improved_resume", "")
        item_ids = improve_body.get("item_ids")

        with open(docx_path, "rb") as f:
            form_data = {"improved_resume": improved_resume_text}
            if item_ids is not None:
                form_data["item_ids"] = json.dumps(item_ids)
            resp = session.post(
                f"{base_url}/api/admin/improve/docx",
                files={"original_file": (docx_path.name, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                data=form_data,
                timeout=60,
            )
        result["timings_sec"]["improve_docx"] = round(time.time() - t0, 2)

        if resp.status_code != 200:
            result["status"] = "FAILED_IMPROVE_DOCX"
            result["error"] = f"HTTP {resp.status_code}: {resp.text[:500]}"
            log(logfile, f"[{docx_path.name}] improve/docx FAILED: {result['error']}")
            return result

        improved_docx_path = out_dir / "improved.docx"
        improved_docx_path.write_bytes(resp.content)
        result["steps"]["improve_docx"] = "OK"
        result["improved_docx_size"] = len(resp.content)
        result["original_docx_size"] = docx_path.stat().st_size

        log(logfile, f"[{docx_path.name}] improve/docx OK ({result['timings_sec']['improve_docx']}s), "
                      f"size {result['original_docx_size']} -> {result['improved_docx_size']} bytes")

        # -------------------------------------------------------------
        # Пост-проверка: сравнение текста original vs improved
        # -------------------------------------------------------------
        try:
            from docx import Document
            orig_doc = Document(str(original_copy))
            impr_doc = Document(str(improved_docx_path))
            orig_paras = [p.text for p in orig_doc.paragraphs if p.text.strip()]
            impr_paras = [p.text for p in impr_doc.paragraphs if p.text.strip()]
            changed = sum(1 for o, i in zip(orig_paras, impr_paras) if o != i)
            result["paragraph_count_original"] = len(orig_paras)
            result["paragraph_count_improved"] = len(impr_paras)
            result["changed_paragraphs"] = changed
            result["table_count_original"] = len(orig_doc.tables)
            result["table_count_improved"] = len(impr_doc.tables)
            log(logfile, f"[{docx_path.name}] Сравнение: {changed}/{len(orig_paras)} параграфов изменено, "
                          f"таблицы {len(orig_doc.tables)}=={len(impr_doc.tables)}")
        except Exception as cmp_err:
            log(logfile, f"[{docx_path.name}] Предупреждение: не удалось сравнить DOCX ({cmp_err})")

        result["status"] = "SUCCESS"

    except Exception as e:
        result["status"] = "EXCEPTION"
        result["error"] = f"{type(e).__name__}: {e}"
        log(logfile, f"[{docx_path.name}] EXCEPTION: {result['error']}")
        log(logfile, traceback.format_exc())

    return result


def main():
    parser = argparse.ArgumentParser(description="Batch runner for ResumeAI regression testing")
    parser.add_argument("--input", required=True, help="Папка с исходными .docx файлами")
    parser.add_argument("--output", required=True, help="Папка для результатов")
    parser.add_argument("--base-url", default="http://127.0.0.1:5000", help="Базовый URL сервера")
    parser.add_argument("--admin-user", default=None,
                         help="Имя/email администратора (или переменная окружения ADMIN_EMAIL)")
    parser.add_argument("--admin-pass", default=None,
                         help="Пароль администратора (или переменная окружения ADMIN_PASSWORD)")
    args = parser.parse_args()

    args.admin_user = args.admin_user or os.environ.get("ADMIN_EMAIL")
    args.admin_pass = args.admin_pass or os.environ.get("ADMIN_PASSWORD")
    if not args.admin_user or not args.admin_pass:
        print("ERROR: укажите учётные данные администратора через --admin-user/--admin-pass "
              "или переменные окружения ADMIN_EMAIL/ADMIN_PASSWORD")
        sys.exit(1)

    input_dir = Path(args.input)
    output_dir = Path(args.output)
    output_dir.mkdir(parents=True, exist_ok=True)

    docx_files = sorted(input_dir.glob("*.docx"))
    if not docx_files:
        print(f"ERROR: в {input_dir} не найдено ни одного .docx файла")
        sys.exit(1)

    main_log_path = output_dir / "batch_log.txt"
    main_logfile = open(main_log_path, "w", encoding="utf-8")

    log(main_logfile, f"=== Batch Runner запущен ===")
    log(main_logfile, f"Input:  {input_dir.resolve()}")
    log(main_logfile, f"Output: {output_dir.resolve()}")
    log(main_logfile, f"Base URL: {args.base_url}")
    log(main_logfile, f"Найдено файлов: {len(docx_files)}")
    for f in docx_files:
        log(main_logfile, f"  - {f.name}")

    session = requests.Session()

    try:
        login(session, args.base_url, args.admin_user, args.admin_pass, main_logfile)
    except Exception as e:
        log(main_logfile, f"КРИТИЧЕСКАЯ ОШИБКА: не удалось залогиниться: {e}")
        log(main_logfile, "Убедитесь что сервер запущен: python run.py")
        main_logfile.close()
        sys.exit(1)

    results = []
    t_batch_start = time.time()

    for docx_path in docx_files:
        file_stem = docx_path.stem
        file_out_dir = output_dir / file_stem
        file_log_path = file_out_dir / "log.txt"
        file_out_dir.mkdir(parents=True, exist_ok=True)

        with open(file_log_path, "w", encoding="utf-8") as file_logfile:
            log(main_logfile, f"--- Обработка {docx_path.name} ---")
            result = run_single_resume(session, args.base_url, docx_path, file_out_dir, file_logfile)
            results.append(result)
            log(main_logfile, f"--- {docx_path.name}: {result['status']} ---")

    total_time = round(time.time() - t_batch_start, 2)

    # Сводка
    summary = {
        "total_files": len(docx_files),
        "success": sum(1 for r in results if r["status"] == "SUCCESS"),
        "failed": sum(1 for r in results if r["status"] != "SUCCESS"),
        "total_time_sec": total_time,
        "results": results,
    }

    (output_dir / "batch_summary.json").write_text(
        json.dumps(summary, indent=2, ensure_ascii=False), encoding="utf-8"
    )

    # Человекочитаемая сводка
    lines = []
    lines.append("=" * 70)
    lines.append("BATCH RUNNER — СВОДНЫЙ ОТЧЁТ")
    lines.append("=" * 70)
    lines.append(f"Всего файлов:   {summary['total_files']}")
    lines.append(f"Успешно:        {summary['success']}")
    lines.append(f"С ошибками:     {summary['failed']}")
    lines.append(f"Общее время:    {total_time}s")
    lines.append("")
    lines.append(f"{'Файл':<28} {'Статус':<20} {'Изменено пар.':<15} {'Токены':<10}")
    lines.append("-" * 70)
    for r in results:
        changed = f"{r.get('changed_paragraphs', '?')}/{r.get('paragraph_count_original', '?')}"
        tokens = str(r.get("tokens_used", "-"))
        lines.append(f"{r['file']:<28} {r['status']:<20} {changed:<15} {tokens:<10}")
        if r.get("error"):
            lines.append(f"    ERROR: {r['error'][:100]}")
        if r.get("quality_summary"):
            qs = r["quality_summary"]
            lines.append(f"    quality: accepted={qs.get('accepted')} kept_original={qs.get('kept_original')} "
                          f"rejected_facts={qs.get('rejected_facts')} needs_retry={qs.get('needs_retry')} "
                          f"avg_sim={qs.get('avg_similarity')}")
    lines.append("=" * 70)

    summary_text = "\n".join(lines)
    (output_dir / "batch_summary.txt").write_text(summary_text, encoding="utf-8")

    log(main_logfile, "")
    log(main_logfile, summary_text)
    main_logfile.close()

    print()
    print(summary_text)
    print(f"\nПолные результаты сохранены в: {output_dir.resolve()}")


if __name__ == "__main__":
    main()
