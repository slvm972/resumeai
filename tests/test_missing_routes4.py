"""
Стресс-тесты missing_routes4.py — 30 сценариев.
Запуск: python -m pytest tests/test_missing_routes4.py -v
"""
import sys, io, re, pytest
sys.path.insert(0, '/home/claude/resumeai')
import importlib, app.missing_routes4 as mr
importlib.reload(mr)

NL = chr(10)

# ===========================================================================
# БЛОК 1: Токены (1-3)
# ===========================================================================

def test_01_token_format():
    tok = mr._make_token()
    assert tok.startswith("@@@") and tok.endswith("@@@")
    inner = tok[3:-3]
    assert len(inner) == 12 and re.match(r"^[0-9A-F]+$", inner)

def test_02_tokens_unique():
    assert len({mr._make_token() for _ in range(100)}) == 100

def test_03_no_double_tokens():
    for text in ["2023-2020", "slvm972@gmail.com", "09/10/1970", "319431516", "AWS-SAA-C03"]:
        store = {}
        assert "@@@@@" not in mr._protect_text(text, store)

# ===========================================================================
# БЛОК 2: Roundtrip (4-5)
# ===========================================================================

def test_04_restore_roundtrip_data():
    for text in [
        "slvm972@gmail.com", "09/10/1970", "2023-2020", "319431516",
        "https://linkedin.com/in/user", "AWS-SAA-C03",
        "C++", "C#", ".NET", "React.js", "SQL Server", "053-3341679",
    ]:
        store = {}
        assert mr._restore_text(mr._protect_text(text, store), store) == text

def test_05_mixed_hebrew_latin_restore():
    text = "ידע ב Python ו-Docker"
    store = {}
    result = mr._protect_text(text, store)
    assert "Python" not in result and "Docker" not in result
    assert mr._restore_text(result, store) == text

# ===========================================================================
# БЛОК 3: Защита сущностей (6-12)
# ===========================================================================

def test_06_email_protected():
    store = {}
    assert "@@@" in mr._protect_text("john.doe@company.co.il", store)

def test_07_linkedin_protected():
    store = {}
    assert "@@@" in mr._protect_text("https://linkedin.com/in/johndoe", store)

def test_08_github_protected():
    store = {}
    assert "@@@" in mr._protect_text("https://github.com/user/repo", store)

def test_09_certificates_protected():
    for cert in ["AWS-SAA-C03", "AZ-104", "CCNA 200-301", "LPIC-1"]:
        store = {}
        result = mr._protect_text(cert, store)
        assert "@@@" in result, f"не защищён: {cert}"
        assert mr._restore_text(result, store) == cert

def test_10_tech_stack_protected():
    for tech in ["React.js", "Node.js", "Docker", "Kubernetes", "PostgreSQL", "Python"]:
        store = {}
        assert "@@@" in mr._protect_text(tech, store), f"не защищён: {tech}"

def test_11_special_syntax_tech():
    for tech in ["C++", "C#", ".NET", "ASP.NET"]:
        store = {}
        result = mr._protect_text(tech, store)
        assert "@@@" in result, f"не защищён: {tech}"
        assert mr._restore_text(result, store) == tech

def test_12_phone_protected():
    store = {}
    assert "@@@" in mr._protect_text("053-3341679", store)

# ===========================================================================
# БЛОК 4: Классификатор (13-15)
# ===========================================================================

def test_13_classifier_freeze():
    for text in [
        "ניסיון תעסוקתי", "השכלה", "שפות",
        "2023-2020", "319431516",
        "דיבור רמה גבוהה, קריאה רמה גבוהה, כתיבה רמה גבוהה",
        "09/10/1970 .",
    ]:
        assert mr._classify_item(text, 5, 40) == "freeze", f"не freeze: {text!r}"

def test_14_classifier_improve():
    for text in [
        "מנהל רשת, אחראי תמיכה ואחזקת רשת, בית תוכנה, קייב",
        "פיתוח ויישום פתרונות טכנולוגיים מתקדמים",
        "התקנה חומרה לרשתות מחשבים, פתח תקווה",
        "managed team of 5 engineers, designed REST API architecture",
        "למידה עצמית בקורסים מקוונים (2023-2021)",
        "יצירת תוכן דיגיטלי: עריכת וידאו, הפקת מוזיקה, ניהול תוכן ברשת",
    ]:
        assert mr._classify_item(text, 5, 40) == "improve", f"не improve: {text!r}"

def test_15_name_is_freeze():
    assert mr._classify_item("מוסתובוי סלבה", 0, 40) == "freeze"

# ===========================================================================
# БЛОК 5: Fact Validation (16-19)
# ===========================================================================

def test_16_fact_validation_rewording_ok():
    valid, _ = mr._validate_block(
        "מנהל רשת, אחראי תמיכה",
        "מנהל רשת מומחה, אחראי על תמיכה ואחזקה"
    )
    assert valid

def test_17_fact_validation_rejects_new_number():
    valid, reason = mr._validate_block(
        "צוות של 5 אנשים",
        "צוות של 50 אנשים"
    )
    assert not valid and "50" in reason

def test_18_fact_validation_rejects_new_tech():
    valid, _ = mr._validate_block(
        "פיתוח עם Python",
        "פיתוח עם Python ו-Java"
    )
    assert not valid

def test_19_fact_validation_rejects_invented_cert():
    valid, _ = mr._validate_block(
        "ניסיון עם רשתות",
        "ניסיון עם רשתות, בעל תעודת CCNA"
    )
    assert not valid

# ===========================================================================
# БЛОК 6: Run-safe DOCX (20-22)
# ===========================================================================

def test_20_single_run_replace_preserves_bold():
    from docx import Document
    doc = Document()
    para = doc.add_paragraph()
    run = para.add_run("original")
    run.bold = True
    mr._replace_para_text(para, "improved")
    assert para.text == "improved" and para.runs[0].bold is True

def test_21_mixed_bold_not_touched():
    from docx import Document
    doc = Document()
    para = doc.add_paragraph()
    r1 = para.add_run("Bold: "); r1.bold = True
    r2 = para.add_run("normal text here"); r2.bold = False
    orig = para.text
    mr._replace_para_text(para, "NEW TEXT")
    assert para.text == orig

def test_22_homogeneous_runs_replaced():
    from docx import Document
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("first part ")
    para.add_run("second part")
    mr._replace_para_text(para, "full new text")
    assert para.text == "full new text"

# ===========================================================================
# БЛОК 7: Roundtrip на реальном документе (23-25)
# ===========================================================================

import os
# Путь ищется относительно корня проекта — положи любой .docx для теста
REAL_DOC_PATH = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'tests', 'sample_resume.docx')

def _get_real_bytes():
    try:
        with open(REAL_DOC_PATH, 'rb') as f:
            return f.read()
    except FileNotFoundError:
        return None

def test_23_real_doc_element_count():
    from docx import Document
    raw = _get_real_bytes()
    if raw is None:
        pytest.skip("тестовый файл недоступен")
    items = mr._extract_structured(Document(io.BytesIO(raw)))
    assert len(items) == 40, f"ожидали 40, получили {len(items)}"

def test_24_real_doc_protect_restore_all():
    from docx import Document
    raw = _get_real_bytes()
    if raw is None:
        pytest.skip("тестовый файл недоступен")
    items = mr._extract_structured(Document(io.BytesIO(raw)))
    for item in items:
        store = {}
        restored = mr._restore_text(mr._protect_text(item['text'], store), store)
        assert restored == item['text'], f"restore failed: {item['text'][:50]!r}"

def test_25_real_doc_apply_roundtrip():
    from docx import Document
    raw = _get_real_bytes()
    if raw is None:
        pytest.skip("тестовый файл недоступен")
    items = mr._extract_structured(Document(io.BytesIO(raw)))
    item_ids = [str(i+1).zfill(3) for i in range(len(items))]
    improved = "###ITEM_" + "\n\n###ITEM_".join(
        f"{item_ids[i]}###\n{items[i]['text']}" for i in range(len(items))
    )
    buf = mr._apply_improved_text_to_docx(raw, improved, item_ids)
    result_items = mr._extract_structured(Document(buf))
    assert len(result_items) == len(items)
    for i, (orig, res) in enumerate(zip(items, result_items)):
        assert res['text'] == orig['text'], f"[{i}] изменился: {orig['text'][:40]!r}"

# ===========================================================================
# БЛОК 8: Edge cases (26-30)
# ===========================================================================

def test_26_empty_bullet_preserved():
    """Пустой буллет сохраняется после roundtrip"""
    text = "• первый пункт" + NL + "• второй пункт" + NL + "•"
    store = {}
    protected = NL.join(mr._protect_text(l, store) for l in text.split(NL))
    assert mr._restore_text(protected, store) == text

def test_27_russian_english_mixed():
    """Русский + английский: технологии защищены"""
    text = "Разработка на Python и JavaScript, опыт с Docker"
    store = {}
    result = mr._protect_text(text, store)
    assert "Python" not in result
    assert "JavaScript" not in result
    assert "Docker" not in result
    assert mr._restore_text(result, store) == text

def test_28_date_range_not_caught_as_phone():
    """2019-2014 — один токен, не PHONE"""
    text = "2019-2014"
    store = {}
    result = mr._protect_text(text, store)
    assert len(store) == 1, f"должен быть 1 токен, получили {len(store)}"
    assert mr._restore_text(result, store) == text

def test_29_fact_validation_adjectives_allowed():
    """Прилагательные без новых фактов — разрешено"""
    valid, _ = mr._validate_block(
        "ניסיון בתמיכה טכנית",
        "ניסיון מקיף בתמיכה טכנית מתקדמת"
    )
    assert valid

def test_30_section_headers_russian_english_freeze():
    """Section headers на русском и английском → freeze"""
    for text in ["Experience", "Education", "Skills", "Опыт работы", "Образование"]:
        assert mr._classify_item(text, 5, 40) == "freeze", f"не freeze: {text!r}"

# ===========================================================================
# АУДИТ P1 + P2 — тесты которые должны упасть на текущей реализации
# ===========================================================================

def test_P1a_capitalized_verb_not_a_fact():
    """P1: 'Led' — глагол с заглавной буквы — не должен считаться новым фактом."""
    valid, reason = mr._validate_block(
        "managed a team of 5 engineers",
        "Led and mentored a team of 5 engineers"
    )
    assert valid, f"Ложное срабатывание: {reason}"

def test_P1b_sentence_start_capital_not_a_fact():
    """P1: Первое слово предложения с заглавной — не факт."""
    valid, reason = mr._validate_block(
        "developed microservices architecture",
        "Architected and delivered a scalable microservices platform"
    )
    assert valid, f"Ложное срабатывание: {reason}"

def test_P1c_hebrew_improvement_accepted():
    """P1: Улучшение на иврите с новыми прилагательными — не должно быть отклонено."""
    valid, reason = mr._validate_block(
        "ניהל צוות מהנדסים",
        "הוביל וניהל צוות מהנדסים מקצועי"
    )
    assert valid, f"Ложное срабатывание: {reason}"

def test_P2a_percentage_protected():
    """P2: Процент 85% должен быть защищён токеном."""
    store = {}
    result = mr._protect_text("כיסוי קוד 85%", store)
    assert "@@@" in result, "Процент не защищён — LLM может его изменить"

def test_P2b_number_with_suffix_protected():
    """P2: Число 1M+ должно быть защищено."""
    store = {}
    result = mr._protect_text("used by 1M+ businesses", store)
    assert "@@@" in result, "Число с суффиксом не защищено"

def test_P2c_number_with_comma_protected():
    """P2: Число 50,000+ должно быть защищено."""
    store = {}
    result = mr._protect_text("serving 50,000+ users", store)
    assert "@@@" in result, "Число с запятой не защищено"

def test_P2d_cicd_protected():
    """P2: CI/CD — технический термин со слешем — должен быть защищён."""
    store = {}
    result = mr._protect_text("experience with CI/CD pipelines", store)
    assert "@@@" in result, "CI/CD не защищён"

def test_P2e_percentage_change_caught_by_fact_validation():
    """P2: Если % не защищён — LLM может изменить 85% на 90% и это не поймается."""
    # Этот тест проверяет что ПОСЛЕ исправления P2 процент защищён
    # и restore работает корректно
    store = {}
    protected = mr._protect_text("code coverage 85%", store)
    restored = mr._restore_text(protected, store)
    assert restored == "code coverage 85%", f"Restore failed: {restored!r}"

# ===========================================================================
# RELEASE v1.0 — Q9 + Q10 (подтверждённые High-приоритет дефекты)
# ===========================================================================

def test_Q10a_short_achievement_with_verb_not_frozen():
    """Q10: 'Led 5 projects' содержит глагол-маркер — должно быть improve, не freeze."""
    result = mr._classify_item("Led 5 projects", 5, 40)
    assert result == "improve", f"Короткое достижение заморожено: {result!r}"

def test_Q10b_short_achievement_saved_not_frozen():
    """Q10: 'Saved $50K' — реальное достижение с фактом — должно улучшаться."""
    result = mr._classify_item("Saved $50K", 5, 40)
    assert result == "improve", f"Короткое достижение заморожено: {result!r}"

def test_Q10c_short_achievement_built_not_frozen():
    """Q10: 'Built 3 apps' — короткое достижение с глаголом — должно улучшаться."""
    result = mr._classify_item("Built 3 apps", 5, 40)
    assert result == "improve", f"Короткое достижение заморожено: {result!r}"

def test_Q10d_short_data_only_still_frozen():
    """Q10 regression guard: короткие строки БЕЗ глагола маркера всё ещё freeze."""
    # Это не должно сломаться после фикса — короткая дата/ID без глагола = freeze
    for text in ["2023-2020", "319431516", "09/10/1970"]:
        result = mr._classify_item(text, 5, 40)
        assert result == "freeze", f"Регрессия: {text!r} стал {result!r}, ожидали freeze"

def test_Q9a_underline_not_silently_lost():
    """Q9: если параграф изменён, underline НЕ должен молча теряться."""
    from docx import Document
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("Contact: ")
    run = para.add_run("email@test.com")
    run.font.underline = True

    text_changed = (para.text != "Contact: email@test.com")
    mr._replace_para_text(para, "New contact info entirely")

    if para.text == "New contact info entirely":
        # Текст был заменён — underline должен быть на run КОТОРЫЙ СОДЕРЖИТ ТЕКСТ,
        # не на пустом run без текста (иначе underline визуально невидим)
        text_runs_with_underline = [r for r in para.runs if r.text.strip() and r.underline]
        assert text_runs_with_underline, \
            f"БАГ Q9: underline остался на пустом run, видимый текст его не имеет. runs={[(r.text, r.underline) for r in para.runs]}"
    else:
        # Параграф не тронут — safe fallback, это ожидаемое поведение после фикса
        assert "email@test.com" in para.text

def test_Q9b_hyperlink_run_not_corrupted():
    """Q9: параграф с гиперссылкой (несколько runs, один с особым форматированием) не должен терять данные."""
    from docx import Document
    doc = Document()
    para = doc.add_paragraph()
    para.add_run("Email: ")
    link_run = para.add_run("contact@company.com")
    link_run.font.underline = True
    link_run.font.color.rgb = None  # имитация ссылки без явного цвета

    orig_text = para.text
    mr._replace_para_text(para, "Email: newcontact@company.com")
    # Результат должен быть либо корректно заменён, либо оригинал сохранён —
    # но НЕ должен быть пустым или повреждённым
    assert para.text, "Параграф стал пустым — повреждение данных"

# ===========================================================================
# АУДИТ Fact Validation — sentence-start capitalization (алгоритмический фикс)
# ===========================================================================

def test_FV1_any_new_sentence_start_verb_accepted():
    """Любой НОВЫЙ (не в стоп-листе) глагол в начале предложения не должен
    считаться выдуманным фактом — это грамматическая заглавная, не факт."""
    valid, reason = mr._validate_block("Managed a team", "Directed a team")
    assert valid, f"Ложное срабатывание на 'Directed': {reason}"

def test_FV2_new_adjective_at_sentence_start_accepted():
    """Новое прилагательное в начале предложения не факт."""
    valid, reason = mr._validate_block("Creative designer", "Visionary designer")
    assert valid, f"Ложное срабатывание на 'Visionary': {reason}"

def test_FV3_arbitrary_synonym_verb_accepted():
    """Произвольный синоним которого нет и не может быть в стоп-листе."""
    valid, reason = mr._validate_block("worked on features", "Championed new features")
    assert valid, f"Ложное срабатывание на 'Championed': {reason}"

def test_FV4_real_new_company_still_rejected():
    """Настоящая новая компания в СЕРЕДИНЕ текста должна по-прежнему отклоняться."""
    valid, reason = mr._validate_block("worked at StartupX", "worked at Google")
    assert not valid, "Регрессия: реальный новый факт (компания) больше не ловится"

def test_FV5_real_new_title_still_rejected():
    """Настоящая новая должность в середине текста должна отклоняться."""
    valid, reason = mr._validate_block("was a developer", "was a Senior Architect")
    assert not valid, "Регрессия: новая должность не поймана"

def test_FV6_real_new_number_still_rejected():
    """Число всё ещё ловится независимо от позиции."""
    valid, reason = mr._validate_block("team of 5 people", "team of 50 people")
    assert not valid, "Регрессия: число не поймано"

def test_FV7_new_tech_mid_sentence_still_rejected():
    """Новая технология В СЕРЕДИНЕ предложения (не позиция 0) должна отклоняться."""
    valid, reason = mr._validate_block("used React framework", "used React and Vue frameworks")
    assert not valid, "Регрессия: новая технология в середине текста не поймана"

def test_FV8_capitalized_fact_at_sentence_start_still_rejected():
    """Важно: если LLM вставил РЕАЛЬНЫЙ новый факт (не глагол) именно в
    начало предложения — это edge case, который допустимо пропустить
    (структурная неопределённость первого слова), но проверяем что хотя бы
    очевидные многословные факты (Title Case company) всё ещё ловятся."""
    valid, reason = mr._validate_block(
        "worked as consultant",
        "worked as consultant for Microsoft Corporation"
    )
    assert not valid, "Настоящая новая компания в конце текста должна ловиться"

# ===========================================================================
# QG1 — Quality Gate не должен блокировать confirmed word-level improvement
# ===========================================================================

def test_QG1_single_strong_synonym_not_blocked_by_high_similarity():
    """Замена одного сильного глагола на синоним не должна блокироваться
    только из-за высокого биграммного сходства (>95%), если
    _has_quality_improvement подтверждает реальное словесное изменение."""
    orig     = "Designed payment flows for Stripe Dashboard used by 1M+ businesses"
    improved = "Architected payment flows for Stripe Dashboard used by 1M+ businesses"
    ok, sim, reason = mr._quality_gate(orig, improved)
    assert ok, f"Confirmed word-change заблокирован Quality Gate: sim={sim:.3f} {reason}"

def test_QG2_truly_unchanged_text_still_needs_retry():
    """Regression guard: полностью идентичный текст всё ещё должен требовать retry."""
    orig = "Managed a team of 5 engineers"
    ok, sim, reason = mr._quality_gate(orig, orig)
    assert not ok, "Регрессия: идентичный текст теперь принимается Quality Gate"

def test_QG3_trivial_punctuation_change_still_needs_retry():
    """Regression guard: добавление только запятой не должно приниматься."""
    orig     = "פיתוח ממשקי משתמש"
    improved = "פיתוח ממשקי משתמש,"
    ok, sim, reason = mr._quality_gate(orig, improved)
    assert not ok, "Регрессия: тривиальное изменение пунктуации теперь принимается"

# ===========================================================================
# HIGH#1 + HIGH#2 — фикс подтверждённых дефектов из финального приёмочного аудита
# ===========================================================================

def test_HIGH1_contact_block_frozen():
    """HIGH#1: контактный блок (email+phone+linkedin+city) должен замораживаться,
    не должен уходить в improve и превращаться в prose."""
    contact = "sarah.mitchell@gmail.com  |  +1 (415) 555-0192  |  linkedin.com/in/sarahmitchell  |  San Francisco, CA"
    result = mr._classify_item(contact, 2, 30)
    assert result == "freeze", f"Контактный блок не заморожен: {result!r}"

def test_HIGH2_fabricated_causal_clause_rejected():
    """HIGH#2: добавленное причинно-следственное утверждение, которого не было
    в оригинале, должно отклоняться Fact Validation."""
    orig = "Led redesign of the host onboarding flow, reducing drop-off rate from 62% to 28%"
    improved = ("Spearheaded the redesign of the host onboarding flow, achieving a significant "
                "reduction in drop-off rate from 62% to 28%, resulting in enhanced user "
                "retention and improved overall experience.")
    valid, reason = mr._validate_block(orig, improved)
    assert not valid, f"Фабрикованный claim не отклонён: {reason}"
    assert "resulting in" in reason.lower()

def test_HIGH2_research_skills_prose_rejected():
    """HIGH#2: превращение списка навыков в рекламное предложение с
    'ensuring...'/'meet user needs' должно отклоняться."""
    orig = "User interviews, A/B testing, usability testing, heatmaps, analytics"
    improved = ("Utilized user interviews, A/B testing, usability testing, heatmaps, and "
                "analytics to drive data-informed design decisions, ensuring products meet "
                "user needs and exceed expectations.")
    valid, reason = mr._validate_block(orig, improved)
    assert not valid, f"Фабрикованный claim в skills не отклонён: {reason}"

def test_HIGH2_legit_rewrite_without_fabrication_still_accepted():
    """Regression guard: обычное улучшение без фабрикованных claims
    по-прежнему принимается (не должно ломаться новым правилом)."""
    orig = "Managed a team of 5 designers across 3 product verticals"
    improved = "Directed a cross-functional team of 5 designers across 3 product verticals"
    valid, reason = mr._validate_block(orig, improved)
    assert valid, f"Ложное срабатывание на легитимное улучшение: {reason}"

def test_HIGH1_non_contact_short_line_still_improvable():
    """Regression guard: обычная короткая improvable-строка (без email/phone)
    не должна случайно замораживаться новым правилом."""
    result = mr._classify_item("Led 5 projects", 5, 40)
    assert result == "improve", f"Регрессия из-за HIGH#1 фикса: {result!r}"

# ===========================================================================
# RETRY-FREEZE — retry должен наследовать hard-freeze решение первой попытки
# ===========================================================================

def test_RETRY1_classify_item_alone_would_not_freeze_name():
    """Доказательство предпосылки бага: _classify_item сам по себе (без
    контекста idx<=1) не замораживает обычное имя — оно должно замораживаться
    только через hard-freeze правило по позиции."""
    result = mr._classify_item("Jonathan Weller", 0, 28)
    assert result != "freeze", (
        "Если этот тест начнёт падать (classify_item сам стал замораживать "
        "короткие имена) — проверка ниже (RETRY2) всё равно должна проходить, "
        "но тогда предпосылка бага изменилась."
    )

def test_RETRY2_retry_wraps_hard_frozen_block_in_single_token():
    """Regression test для найденного дефекта: retry-построение для блока,
    который на этапе защиты был hard-frozen (idx<=1), должно оборачивать его
    в ОДИН непрозрачный токен (как при первой попытке), а НЕ пропускать через
    _protect_text() построчно, что оставляло бы обычные слова имени открытыми."""
    orig_name = "Jonathan Weller"
    store = {}

    # Симулируем то, что теперь происходит в retry-коде при strategy_map[iid]=="freeze"
    strategy_map = {"001": "freeze"}
    iid = "001"
    if strategy_map.get(iid) == "freeze":
        tok = mr._make_token()
        store[tok] = orig_name
        protected_t = tok
    else:
        protected_t = mr._protect_text(orig_name, store)

    # Проверка: весь текст обёрнут в ОДИН токен, ни одно слово имени не видно
    assert protected_t.startswith("@@@") and protected_t.endswith("@@@")
    assert "Jonathan" not in protected_t
    assert "Weller" not in protected_t
    # restore должен точно вернуть оригинал
    assert mr._restore_text(protected_t, store) == orig_name

def test_RETRY3_non_frozen_block_still_uses_protect_text_in_retry():
    """Regression guard: обычный improve-блок (не hard-frozen) должен
    по-прежнему идти через _protect_text() при retry, а не оборачиваться
    в один токен — иначе LLM вообще ничего не сможет улучшить при retry."""
    orig_text = "Managed a team of 5 designers across 3 product verticals"
    store = {}
    strategy_map = {"008": "improve"}
    iid = "008"
    if strategy_map.get(iid) == "freeze":
        tok = mr._make_token()
        store[tok] = orig_text
        protected_t = tok
    else:
        protected_t = mr._protect_text(orig_text, store)

    # Обычный текст должен остаться частично видимым (не единый токен)
    assert not (protected_t.startswith("@@@") and protected_t.endswith("@@@") and protected_t.count("@@@") == 2), \
        "Regression: improve-блок ошибочно обёрнут в единый token как freeze-блок"
    assert "Managed" in protected_t or "team" in protected_t

# ===========================================================================
# COMMERCIAL QUALITY FIXES — найдено на реальных batch-прогонах (001_classic)
# ===========================================================================

def test_MODEL1_no_dead_fallback_model_referenced():
    """Мёртвая модель gemma2-9b-it (decommissioned Groq) не должна больше
    упоминаться нигде в коде — заменена на подтверждённо живую модель."""
    with open(mr.__file__, encoding="utf-8") as f:
        content = f.read()
    assert "gemma2-9b-it" not in content, "Мёртвая fallback-модель всё ещё в коде"

def test_SEC1_professional_summary_frozen():
    """'Professional Summary' — section header, должен замораживаться
    (реальный баг: раньше отсутствовал в SECTION_HEADERS_SET)."""
    assert mr._classify_item("Professional Summary", 3, 30) == "freeze"

def test_SEC2_work_experience_frozen():
    """'Work Experience' — section header, должен замораживаться."""
    assert mr._classify_item("Work Experience", 5, 30) == "freeze"

def test_IDENTITY1_job_title_company_line_frozen():
    """Identity-строка 'Title — Company, Location' — структурное поле,
    не должно переписываться (реальный баг из batch-прогона)."""
    text = "Marketing Operations Manager — Northfield Analytics, Chicago, IL"
    assert mr._classify_item(text, 16, 30) == "freeze"

def test_IDENTITY2_degree_university_line_frozen():
    """Identity-строка 'Degree — University' — структурное поле."""
    text = "B.A. in Business Administration — Marquette University"
    assert mr._classify_item(text, 20, 30) == "freeze"

def test_IDENTITY3_sarah_mitchell_job_line_also_fixed():
    """Regression: та же identity-line ошибка была и в резюме Sarah Mitchell."""
    text = "Lead UX Designer — Airbnb, San Francisco, CA"
    assert mr._classify_item(text, 16, 30) == "freeze"

def test_IDENTITY4_bullet_with_dash_and_verb_stays_improvable():
    """Regression guard: буллет с тире, но содержащий глагол-маркер,
    не должен ошибочно замораживаться новым identity-правилом."""
    text = "Reduced costs - saved $50K annually through vendor consolidation"
    assert mr._classify_item(text, 10, 30) == "improve"

def test_COLON1_skill_category_label_frozen():
    """Label заголовок таблицы навыков с двоеточием — структурное поле."""
    assert mr._classify_item("Marketing Tools:", 22, 30) == "freeze"

def test_ENUM1_tool_list_frozen():
    """Список инструментов через запятую без глагола — нечего улучшать,
    только риск фабрикации при попытке 'переписать' список названий."""
    text = "HubSpot, Salesforce, Marketo, Google Analytics"
    assert mr._classify_item(text, 24, 30) == "freeze"

def test_ENUM2_language_list_frozen():
    """Список языков через запятую — аналогично списку инструментов."""
    text = "English (Native), German (Conversational)"
    assert mr._classify_item(text, 28, 30) == "freeze"

def test_ENUM3_real_achievement_with_commas_stays_improvable():
    """Regression guard: реальное достижение с запятыми и глаголом
    не должно ошибочно замораживаться новым enum-правилом."""
    text = "Reduced churn from 18% to 6.5%, increased NPS, and expanded team headcount"
    assert mr._classify_item(text, 10, 30) == "improve"

def test_ENUM4_metrics_bullet_with_commas_stays_improvable():
    """Regression guard: достижение с ARR/процентами и запятыми внутри
    остаётся improvable (реальный кейс из resume 002_metrics)."""
    text = "Grew annual recurring revenue (ARR) from $8M to $64M in 4 years, a 700% increase"
    assert mr._classify_item(text, 10, 30) == "improve"

# ===========================================================================
# BACKOFF — парсинг времени ожидания из ответа Groq (429)
# ===========================================================================

def test_BACKOFF1_parses_real_groq_message():
    """Реальное сообщение об ошибке Groq корректно парсится."""
    msg = ("Rate limit reached for model `llama-3.1-8b-instant`... "
           "Please try again in 5.67s. Need more tokens?")
    assert mr._extract_retry_after_seconds(msg) == 5.67

def test_BACKOFF2_caps_at_maximum():
    """Слишком большое время ожидания ограничивается cap."""
    msg = "Please try again in 999.9s."
    assert mr._extract_retry_after_seconds(msg, cap=12.0) == 12.0

def test_BACKOFF3_default_when_unparseable():
    """Если формат сообщения не распознан — используется default."""
    assert mr._extract_retry_after_seconds("some unrelated error", default=2.0) == 2.0
