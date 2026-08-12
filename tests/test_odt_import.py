"""
Тесты для поддержки .odt на вход в _extract_text_from_request()
(app/__init__.py). Проверяем реальным .odt файлом (сконвертированным
из .docx через LibreOffice), не заглушкой — включая многоязычность
(латиница + иврит) и таблицы.

Запуск: python -m pytest tests/test_odt_import.py -v
"""
import os
import sys
import io

sys.path.insert(0, '/home/claude/resumeai')

# Безопасные тестовые значения ДО ЛЮБОГО импорта app/config — тот же паттерн,
# что и в tests/test_credits_pool.py / tests/test_odt_export.py.
os.environ['FLASK_ENV'] = 'testing'
os.environ.setdefault('SECRET_KEY', 'test-secret-key-for-tests-only')
os.environ.setdefault('JWT_SECRET_KEY', 'test-jwt-secret-for-tests-only')
os.environ.setdefault('GROQ_API_KEY', 'test-groq-key-not-used-mocked-out')

import pytest
from docx import Document

from app import create_app, _extract_text_from_request

APP = create_app('testing')

SAMPLE_ODT_PATH = os.path.join(os.path.dirname(__file__), 'sample_import.odt')


def _build_sample_odt(path):
    """
    Собрать тестовый .odt через LibreOffice headless конвертацию из .docx —
    тот же метод, что уже применялся в проекте для получения тестовых .odt.
    Содержит латиницу + иврит + таблицу (как реальное резюме).
    """
    import subprocess
    import tempfile

    tmp_dir = tempfile.mkdtemp()
    docx_path = os.path.join(tmp_dir, 'sample_import.docx')

    doc = Document()
    doc.add_paragraph("John Doe")
    doc.add_paragraph("Software Engineer")
    doc.add_paragraph("ניסיון תעסוקתי")
    doc.add_paragraph("מנהל רשת, אחראי תמיכה ואחזקת רשת, בית תוכנה, קייב")

    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Skills"
    table.cell(0, 1).text = "Python, Docker"
    table.cell(1, 0).text = "כישורים"
    table.cell(1, 1).text = "עברית, אנגלית"

    doc.save(docx_path)

    result = subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "odt",
         "--outdir", tmp_dir, docx_path],
        capture_output=True, text=True, timeout=60,
    )
    converted = os.path.join(tmp_dir, 'sample_import.odt')
    if result.returncode != 0 or not os.path.exists(converted):
        pytest.skip(f"LibreOffice недоступен для конвертации: {result.stderr}")

    import shutil
    shutil.copy(converted, path)


@pytest.fixture(scope="module", autouse=True)
def ensure_sample_odt():
    if not os.path.exists(SAMPLE_ODT_PATH):
        _build_sample_odt(SAMPLE_ODT_PATH)
    yield


def _extract_from_odt_bytes(odt_bytes, filename='sample_import.odt'):
    with APP.test_request_context(
        '/api/admin/analyze', method='POST',
        data={'file': (io.BytesIO(odt_bytes), filename)},
        content_type='multipart/form-data',
    ):
        return _extract_text_from_request()


def test_01_odt_extracts_latin_paragraphs():
    if not os.path.exists(SAMPLE_ODT_PATH):
        pytest.skip("тестовый .odt недоступен")
    with open(SAMPLE_ODT_PATH, 'rb') as f:
        text = _extract_from_odt_bytes(f.read())
    assert "John Doe" in text
    assert "Software Engineer" in text


def test_02_odt_extracts_hebrew_paragraphs():
    if not os.path.exists(SAMPLE_ODT_PATH):
        pytest.skip("тестовый .odt недоступен")
    with open(SAMPLE_ODT_PATH, 'rb') as f:
        text = _extract_from_odt_bytes(f.read())
    assert "ניסיון תעסוקתי" in text
    assert "מנהל רשת" in text


def test_03_odt_extracts_table_cells():
    """Таблицы (резюме часто хранятся в таблицах) должны читаться, как и у DOCX."""
    if not os.path.exists(SAMPLE_ODT_PATH):
        pytest.skip("тестовый .odt недоступен")
    with open(SAMPLE_ODT_PATH, 'rb') as f:
        text = _extract_from_odt_bytes(f.read())
    assert "Python, Docker" in text
    assert "עברית, אנגלית" in text


def test_04_odt_no_duplicate_table_text():
    """
    Регрессия: текст ячеек таблицы не должен попадать дважды (один раз как
    'обычный параграф' верхнего уровня, второй раз как содержимое таблицы) —
    именно эта ошибка эмпирически возникала при наивном обходе всех <text:p>.
    """
    if not os.path.exists(SAMPLE_ODT_PATH):
        pytest.skip("тестовый .odt недоступен")
    with open(SAMPLE_ODT_PATH, 'rb') as f:
        text = _extract_from_odt_bytes(f.read())
    assert text.count("Python, Docker") == 1
    assert text.count("Skills") == 1


def test_05_odt_no_garbage_bytes():
    """
    Регрессия: до фикса .odt проваливался в catch-all decode('utf-8'), что
    давало нечитаемый мусор (сырые ZIP-байты). Проверяем что результат
    выглядит как связный текст, а не бинарные ошибки декодирования.
    """
    if not os.path.exists(SAMPLE_ODT_PATH):
        pytest.skip("тестовый .odt недоступен")
    with open(SAMPLE_ODT_PATH, 'rb') as f:
        text = _extract_from_odt_bytes(f.read())
    assert "\ufffd" not in text  # replacement character — признак мусора


def test_06_corrupted_odt_raises_value_error():
    """Повреждённый .odt (не ZIP) должен давать понятную ValueError, не падать сырым traceback."""
    garbage = b"this is not a valid odt file at all, just plain bytes"
    with pytest.raises(ValueError):
        _extract_from_odt_bytes(garbage, filename='broken.odt')


def test_08_odt_heading_text_not_lost():
    """
    Регрессия: заголовки (Heading 1/2/...) сохраняются в OpenDocument XML
    как отдельный элемент <text:h>, а не <text:p>. Найден реальный баг —
    ранний фильтр (только qn[1]=='p') исключал <text:h>, из-за чего имя
    резюме и названия секций, стилизованные как заголовки, молча пропадали
    из извлечённого текста. Собираем .odt с заголовками через
    add_heading() в исходном .docx и проверяем, что текст присутствует.
    """
    import subprocess
    import tempfile
    import shutil

    tmp_dir = tempfile.mkdtemp()
    docx_path = os.path.join(tmp_dir, 'heading_sample.docx')

    doc = Document()
    doc.add_heading("John Doe Resume", level=1)
    doc.add_paragraph("Software Engineer")
    doc.add_heading("Experience", level=2)
    doc.add_paragraph("Worked at Acme Corp for 5 years.")
    doc.save(docx_path)

    result = subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "odt",
         "--outdir", tmp_dir, docx_path],
        capture_output=True, text=True, timeout=60,
    )
    converted = os.path.join(tmp_dir, 'heading_sample.odt')
    if result.returncode != 0 or not os.path.exists(converted):
        pytest.skip(f"LibreOffice недоступен для конвертации: {result.stderr}")

    with open(converted, 'rb') as f:
        odt_bytes = f.read()

    text = _extract_from_odt_bytes(odt_bytes, filename='heading_sample.odt')

    assert "John Doe Resume" in text, "Заголовок уровня 1 (<text:h>) потерян"
    assert "Experience" in text, "Заголовок уровня 2 (<text:h>) потерян"
    assert "Software Engineer" in text
    assert "Worked at Acme Corp for 5 years." in text

    shutil.rmtree(tmp_dir, ignore_errors=True)


def test_07_docx_branch_still_works_after_odt_addition():
    """Regression guard: добавление ODT-ветки не должно было сломать существующую DOCX-ветку."""
    doc = Document()
    doc.add_paragraph("Regression Check DOCX")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    with APP.test_request_context(
        '/api/admin/analyze', method='POST',
        data={'file': (buf, 'regression.docx')},
        content_type='multipart/form-data',
    ):
        text = _extract_text_from_request()
    assert "Regression Check DOCX" in text
