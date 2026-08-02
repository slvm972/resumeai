"""
Регрессионные тесты для бага с дублированием merge-ячеек таблиц в
_extract_structured() (app/missing_routes4.py).

Причина бага: python-docx возвращает один и тот же физический объект
_Cell на каждой позиции сетки, которую перекрывает merge-ячейка.
Старый dedup-ключ (id(table), ri, ci) был построен на индексах позиции
в цикле, а не на идентичности самой ячейки, поэтому merge-ячейки
дублировались в результате. Исправлено на id(cell._tc) — идентичность
нижележащего XML-элемента <w:tc>, которая по природе формата OOXML
одна и та же для всех позиций сетки, охваченных merge.

Запуск: python -m pytest tests/test_merged_cells_regression.py -v
"""
import io
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))

from docx import Document
from app.missing_routes4 import _extract_structured, _apply_improved_text_to_docx


REAL_DOC_PATH = os.path.join(os.path.dirname(__file__), "sample_resume.docx")


def test_A_horizontal_merge_not_duplicated():
    """Горизонтальный merge (1x2 -> 1 ячейка): текст должен появиться
    ровно один раз, а не дублироваться на обе исходные позиции сетки."""
    doc = Document()
    table = doc.add_table(rows=1, cols=2)

    cell_a = table.cell(0, 0)
    cell_b = table.cell(0, 1)
    cell_a.merge(cell_b)
    cell_a.text = "Section Header"

    items = _extract_structured(doc)
    matching = [it for it in items if it["text"] == "Section Header"]

    assert len(matching) == 1, (
        f"Ожидали ровно 1 item 'Section Header' после горизонтального merge, "
        f"получили {len(matching)}. Все items: {[it['text'] for it in items]}"
    )


def test_B_vertical_merge_not_duplicated_and_siblings_intact():
    """Вертикальный merge (левая колонка 2x1 -> 1 ячейка) не должен
    дублировать merge-текст и не должен ломать/сдвигать несвязанный
    контент в правой колонке."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)

    left_top = table.cell(0, 0)
    left_bottom = table.cell(1, 0)
    left_top.merge(left_bottom)
    left_top.text = "Merged Vertical"

    table.cell(0, 1).text = "Row1 Right"
    table.cell(1, 1).text = "Row2 Right"

    items = _extract_structured(doc)
    texts = [it["text"] for it in items]

    assert len(items) == 3, (
        f"Ожидали ровно 3 item (1 merge + 2 несвязанных), получили {len(items)}: {texts}"
    )
    assert texts.count("Merged Vertical") == 1, f"Merge-текст продублирован: {texts}"
    assert "Row1 Right" in texts, f"Потерян несвязанный текст Row1 Right: {texts}"
    assert "Row2 Right" in texts, f"Потерян несвязанный текст Row2 Right: {texts}"


def test_C_real_doc_no_merge_identity_roundtrip():
    """Regression guard: реальный документ без merge-ячеек (sample_resume.docx)
    должен проходить identity-roundtrip без единого расхождения — фикс
    дедупликации не должен ничего сломать на обычных (не merge) таблицах."""
    if not os.path.exists(REAL_DOC_PATH):
        import pytest
        pytest.skip("тестовый файл sample_resume.docx недоступен")

    with open(REAL_DOC_PATH, "rb") as f:
        raw = f.read()

    doc_before = Document(io.BytesIO(raw))
    items_before = _extract_structured(doc_before)

    item_ids = [str(i + 1).zfill(3) for i in range(len(items_before))]
    improved_text = "###ITEM_" + "\n\n###ITEM_".join(
        f"{item_ids[i]}###\n{items_before[i]['text']}" for i in range(len(items_before))
    )

    buf = _apply_improved_text_to_docx(raw, improved_text, item_ids)

    doc_after = Document(buf)
    items_after = _extract_structured(doc_after)

    assert len(items_after) == len(items_before), (
        f"Количество items изменилось: было {len(items_before)}, стало {len(items_after)}"
    )
    for i, (before, after) in enumerate(zip(items_before, items_after)):
        assert after["text"] == before["text"], (
            f"[{i}] текст изменился: было {before['text'][:60]!r}, стало {after['text'][:60]!r}"
        )
